#!/usr/bin/env python3
"""Generate comprehensive site analysis report for auspost.com.au"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================
# TITLE PAGE
# ============================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Australia Post (auspost.com.au)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Comprehensive Site Analysis Report', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Migration Scope & Complexity Assessment\n')
run.bold = True
info.add_run('Date: June 2026\n')
info.add_run('Scope: All pages excluding eCommerce hierarchy (/shop)\n')
info.add_run('Source URL: https://auspost.com.au/')

doc.add_page_break()

# ============================================
# TABLE OF CONTENTS
# ============================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Templates Inventory',
    '2. Blocks / Components Catalog',
    '3. Page Counts by Template',
    '4. Integrations Analysis',
    '5. Complex Use Cases & Observations',
    '6. Migration Estimates',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================
# SECTION 1: TEMPLATES INVENTORY
# ============================================
doc.add_heading('1. Templates Inventory', level=1)
doc.add_paragraph(
    'The following unique page templates have been identified across auspost.com.au '
    '(excluding eCommerce/shop pages). Templates are categorized based on their '
    'structural layout, content patterns, and functional purpose.'
)

# Templates table
templates_data = [
    ['Template Name', 'Complexity', 'Reasoning', 'Reference URL Examples'],
    [
        'Homepage',
        'High',
        'Hero carousel, tracking widget, quick-links grid, promotional cards carousel, delivery info section, shop promo, I-want-to links, multiple interactive components',
        'https://auspost.com.au/'
    ],
    [
        'Category Landing / Hub Page',
        'Medium',
        'Section hero with breadcrumbs, H1 + description, content card grid (image + text + CTA), FAQ accordion, footnotes. Used for primary service categories.',
        'https://auspost.com.au/sending/delivery-speeds-and-coverage\nhttps://auspost.com.au/disruptions-and-updates\nhttps://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'Service/Product Detail Page',
        'Medium',
        'Hero banner, feature bullet list with icons, step-by-step process cards, embedded location finder (Google Maps), feature cards grid, FAQ accordion, disclaimer/footnotes.',
        'https://auspost.com.au/receiving/parcel-deliveries/parcel-lockers\nhttps://auspost.com.au/money-travel/banking-and-paying-bills/bank-at-post'
    ],
    [
        'Content / Story Page',
        'Medium',
        'Hero image with overlay text, rich body content, embedded YouTube videos, image+text two-column blocks, card grids with CTAs, partnership content, transcript toggles.',
        'https://auspost.com.au/about-us/supporting-communities/mental-health\nhttps://auspost.com.au/about-us/sustainability'
    ],
    [
        'Tool / Calculator Page',
        'High',
        'Interactive form-based tool (currency converter, postage calculator), real-time API integration, dynamic results display, related product cards below.',
        'https://auspost.com.au/currency-converter\nhttps://auspost.com.au/parcels-mail/calculate-postage-delivery-times/'
    ],
    [
        'Help & Support Hub',
        'High',
        'Custom layout with category cards (icon + title + description), dedicated search bar, multi-channel contact section (online forms, chat widget, phone), live chat integration.',
        'https://auspost.com.au/help-and-support'
    ],
    [
        'Location Finder / Map Page',
        'High',
        'Google Maps integration, address autocomplete, location type filters, service filters, proximity search, state browsing, dynamic results rendering.',
        'https://auspost.com.au/locate'
    ],
    [
        'Location Directory / Index Page',
        'Low',
        'Simple alphabetical listing of suburbs by state, dropdown filters for state and letter, link list to individual location pages. Highly templatized.',
        'https://auspost.com.au/locate/post-office/vic\nhttps://auspost.com.au/locate/post-office/nsw'
    ],
    [
        'Enterprise & Gov Landing',
        'Medium',
        'Custom hero with enterprise messaging, capability cards, solution cards, insights/reports section, customer story links. Distinct from Personal/Business pages.',
        'https://auspost.com.au/enterprise-gov'
    ],
    [
        'Business Landing Page',
        'Medium',
        'Business-focused hero, service category cards, shipping solutions overview, integration CTAs (MyPost Business, eParcel), business-specific content blocks.',
        'https://auspost.com.au/business'
    ],
    [
        'Information / Policy Page',
        'Low',
        'Simple long-form text content with headings, tables for data (e.g., postcode tables, pricing), embedded links, footnotes. Minimal interactive elements.',
        'https://auspost.com.au/sending/delivery-speeds-and-coverage\nhttps://auspost.com.au/privacy\nhttps://auspost.com.au/terms-conditions'
    ],
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
for i, header in enumerate(templates_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in templates_data[1:]:
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()
doc.add_paragraph(
    'Note: The site is built on Adobe Experience Manager (AEM) with extensive use of '
    'client-side JavaScript for navigation and interactive features. The mega-navigation '
    'is shared across all templates.'
)

doc.add_page_break()

# ============================================
# SECTION 2: BLOCKS / COMPONENTS CATALOG
# ============================================
doc.add_heading('2. Blocks / Components Catalog', level=1)
doc.add_paragraph(
    'The following reusable blocks and components have been identified across the site. '
    'Components are listed with their design variations noted where applicable.'
)

blocks_data = [
    ['Block Name', 'Complexity', 'Description & Behaviour', 'Reference URL(s)'],
    [
        'Global Header / Mega Navigation',
        'High',
        'Multi-level mega menu with 3 audience tabs (Personal, Business, Enterprise & Gov). Collapsible mobile hamburger. Quick links panel (Track, Postcode, etc). Auth0 login button. Coveo-powered search overlay. Responsive breakpoints.',
        'All pages\nhttps://auspost.com.au/'
    ],
    [
        'Global Footer',
        'Medium',
        'Two sections: (1) Collapsible "Quick links" and "Read our blogs" accordions, (2) Global footer with legal links, social media icons (Facebook, X, LinkedIn), Help & support link, Aboriginal acknowledgement statement.',
        'All pages\nhttps://auspost.com.au/'
    ],
    [
        'Hero Banner',
        'Medium',
        'Full-width hero with background image, H1 heading, description paragraph, and CTA button with arrow icon. Variant with colour overlay (red gradient). Used on homepage and landing pages.',
        'https://auspost.com.au/\nhttps://auspost.com.au/about-us/supporting-communities/mental-health'
    ],
    [
        'Tracking Widget',
        'High',
        'Inline search form with text input for tracking number, "Track" button. Integrates with MyPost tracking API. Positioned prominently below hero on homepage.',
        'https://auspost.com.au/'
    ],
    [
        'Quick Links Bar (Icon + Text)',
        'Low',
        'Horizontal scrollable list of icon+text link items. Each item has an icon, label, and chevron arrow. Used for primary service shortcuts.',
        'https://auspost.com.au/ (below tracking widget)'
    ],
    [
        'Promotional Card (Image + Text + CTA)',
        'Medium',
        'Card with optional image, H4 heading, description paragraph, and one or more CTA links. Used in carousels and grids. Variants: with image, without image, multi-CTA, horizontal layout.',
        'https://auspost.com.au/\nhttps://auspost.com.au/currency-converter\nhttps://auspost.com.au/disruptions-and-updates'
    ],
    [
        'Content Card Carousel',
        'Medium',
        'Horizontally scrollable card container showing 2-3 promotional cards at a time with navigation dots/status indicators. Auto-advances.',
        'https://auspost.com.au/ (middle section)'
    ],
    [
        'Section Heading + Link List',
        'Low',
        'H2 heading with description paragraph, followed by a vertical list of text links with arrow icons. Used for "Delivery information" section pattern.',
        'https://auspost.com.au/ (Delivery information section)'
    ],
    [
        'Shop Promo Block (Image Grid + Text)',
        'Medium',
        'Three image links in a row (Postage, Collectables, Gifts) with a right-side text panel containing eyebrow, heading, description and "Shop now" CTA.',
        'https://auspost.com.au/ (Buy online section)'
    ],
    [
        'I Want To (Quick Links Grid)',
        'Low',
        'Vertical list of icon + text links for quick actions. H2 heading "I want to" followed by 10 action links with icons.',
        'https://auspost.com.au/ (bottom section)'
    ],
    [
        'Breadcrumbs',
        'Low',
        'Horizontal breadcrumb navigation showing page hierarchy (Personal > Category > Page). Separator arrows between items.',
        'All content pages\nhttps://auspost.com.au/sending/delivery-speeds-and-coverage'
    ],
    [
        'Data Table',
        'Medium',
        'Responsive table with header row, service names as row headers with links, data cells. Used for delivery speeds, postcode ranges, pricing.',
        'https://auspost.com.au/sending/delivery-speeds-and-coverage'
    ],
    [
        'Feature List (Icon + Text)',
        'Low',
        'Vertical bulleted list where each item has a checkmark/tick icon followed by descriptive text. Used to highlight service benefits.',
        'https://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'Step-by-Step Cards',
        'Medium',
        'Numbered cards (1, 2, 3) with H4 heading and description. Used to explain processes (e.g., "How to use a Parcel Locker"). Cards arranged in a grid.',
        'https://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'Location Finder (Google Maps)',
        'High',
        'Embedded Google Maps with Places Autocomplete search, location type dropdown filter, service filter, geolocation button, search results overlay. Complex third-party integration.',
        'https://auspost.com.au/locate\nhttps://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'FAQ Accordion',
        'Medium',
        'Collapsible FAQ section with H2 heading "Frequently asked questions". Individual questions expand to reveal answers. Dynamically loaded content. "View more FAQs" link.',
        'https://auspost.com.au/sending/delivery-speeds-and-coverage\nhttps://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'Video Embed (YouTube)',
        'Medium',
        'Embedded YouTube iframe player with transcript toggle button. Used for partnership/story content. Multiple videos can appear in a grid layout.',
        'https://auspost.com.au/about-us/supporting-communities/mental-health'
    ],
    [
        'CTA Banner (Full Width)',
        'Low',
        'Full-width coloured banner with heading, description text, and primary CTA button. Used as call-to-action section divider.',
        'https://auspost.com.au/receiving/parcel-deliveries/parcel-lockers\nhttps://auspost.com.au/sending/delivery-speeds-and-coverage'
    ],
    [
        'Currency Converter Tool',
        'High',
        'Interactive form with amount input (AUD), destination currency searchable dropdown, and "Go" button. Real-time exchange rate API integration. Results display.',
        'https://auspost.com.au/currency-converter'
    ],
    [
        'Help Category Cards',
        'Medium',
        'Large clickable cards with icon, H3 title, and description. Used on Help & Support page to route users to correct support channel. 4-card grid layout.',
        'https://auspost.com.au/help-and-support'
    ],
    [
        'Contact Channel Block',
        'Medium',
        'Three-column layout with icon, heading (Enquire online / Chat / Call us), description, and action link/button. Chat has "Ask a question" button. Call has expandable phone numbers.',
        'https://auspost.com.au/help-and-support'
    ],
    [
        'Search Bar (Coveo)',
        'High',
        'Full-width search input with placeholder text and search button. Powered by Coveo search platform. Returns AI-assisted search results. Used on Help and global search.',
        'https://auspost.com.au/help-and-support\nAll pages (header search)'
    ],
    [
        'App Download Banner',
        'Low',
        'Promotional block with bulleted features list, App Store and Google Play badge links. Used to promote AusPost mobile app.',
        'https://auspost.com.au/disruptions-and-updates'
    ],
    [
        'Disclaimer / Footnotes',
        'Low',
        'Text block with superscript number references and corresponding footnote paragraphs with legal/disclaimer text. Common at page bottom.',
        'https://auspost.com.au/currency-converter\nhttps://auspost.com.au/receiving/parcel-deliveries/parcel-lockers'
    ],
    [
        'Alert / Notice Banner',
        'Low',
        'Inline notification bar with icon, text message, and optional dismiss button. Used for cookie consent and service disruption alerts at page top.',
        'https://auspost.com.au/ (cookie banner)\nhttps://auspost.com.au/locate'
    ],
    [
        'Suburb Index / Directory Listing',
        'Low',
        'Alphabetical listing component with state dropdown, letter filter tabs, and linked list of suburb names. Used for location directory pages.',
        'https://auspost.com.au/locate/post-office/vic'
    ],
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Medium Shading 1 Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table2.rows[0].cells
for i, header in enumerate(blocks_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in blocks_data[1:]:
    row_cells = table2.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()

# Add screenshots reference
doc.add_heading('Screenshots Reference', level=2)
doc.add_paragraph('The following screenshots have been captured for visual reference:')

screenshots_list = [
    ('Homepage (full page)', '/tmp/playwright/screenshots/homepage.png'),
    ('Content Page - Delivery Speeds & Coverage', '/tmp/playwright/screenshots/content-page-delivery.png'),
    ('Content Page - Mental Health (Story template)', '/tmp/playwright/screenshots/content-page-mental-health.png'),
    ('Service Page - Parcel Lockers', '/tmp/playwright/screenshots/service-page-parcel-lockers.png'),
    ('Tool Page - Currency Converter', '/tmp/playwright/screenshots/tool-page-currency-converter.png'),
    ('Location Finder', '/tmp/playwright/screenshots/locate-page.png'),
    ('Help & Support Hub', '/tmp/playwright/screenshots/help-support-page.png'),
]

for title_text, path in screenshots_list:
    doc.add_paragraph(f'{title_text}', style='List Bullet')
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(4.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'  [Screenshot file: {path}]')
    doc.add_paragraph()

doc.add_page_break()

# ============================================
# SECTION 3: PAGE COUNTS BY TEMPLATE
# ============================================
doc.add_heading('3. Page Counts by Template', level=1)
doc.add_paragraph(
    'Based on sitemap analysis (600+ URLs total, ~377+ excluding eCommerce/shop), '
    'the following distribution has been estimated by template type:'
)

page_counts_data = [
    ['Template Type', 'Est. Pages', 'Auto-Migration', 'Manual Migration', 'Notes'],
    ['Homepage', '1', 'No', 'Yes', 'Unique layout, multiple interactive widgets, heavy customization'],
    ['Category Landing / Hub Page', '~25', 'Partial', 'Partial', 'Standard card layouts auto-migratable; custom widgets need manual work'],
    ['Service/Product Detail Page', '~80', 'Yes (mostly)', 'Some', 'Standardized content model; Google Maps embeds need manual integration'],
    ['Content / Story Page', '~45', 'Yes (mostly)', 'Some', 'Rich text + images auto-migratable; YouTube embeds and interactive elements manual'],
    ['Tool / Calculator Page', '~8', 'No', 'Yes', 'Heavy API integrations, dynamic forms, real-time data - custom development'],
    ['Help & Support Hub', '1', 'No', 'Yes', 'Unique layout with chat, search, and multi-channel contact integration'],
    ['Location Finder / Map Page', '1', 'No', 'Yes', 'Google Maps + Places API + custom location database integration'],
    ['Location Directory / Index', '~16', 'Yes', 'No', '8 states x 2 types (Post Office + Parcel Locker). Templatized, data-driven.'],
    ['Enterprise & Gov Landing', '~35', 'Partial', 'Partial', 'Some pages content-heavy (auto), some have custom interactive elements'],
    ['Business Landing Page', '~55', 'Partial', 'Partial', 'Similar to Enterprise; standard content is auto, integrations are manual'],
    ['Information / Policy Page', '~60', 'Yes', 'No', 'Simple long-form text with tables and links. Highly standardized.'],
    ['Travel Tips & Guides', '~50', 'Yes', 'No', 'Article-style content pages with consistent template structure'],
    ['TOTAL (excl. eCommerce)', '~377', '~200 auto', '~177 manual', ''],
]

table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Medium Shading 1 Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table3.rows[0].cells
for i, header in enumerate(page_counts_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in page_counts_data[1:]:
    row_cells = table3.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()
doc.add_heading('Migration Classification Summary', level=2)
doc.add_paragraph(
    'Automatically Migratable (~53%): Pages with standardized content models, '
    'simple text + image layouts, consistent templates (policy pages, guides, '
    'directory listings, standard content pages)', style='List Bullet'
)
doc.add_paragraph(
    'Requires Manual Migration (~47%): Pages with dynamic functionality, '
    'third-party integrations, custom interactive tools, unique layouts, '
    'and complex JavaScript-driven features', style='List Bullet'
)

doc.add_page_break()

# ============================================
# SECTION 4: INTEGRATIONS ANALYSIS
# ============================================
doc.add_heading('4. Integrations Analysis', level=1)
doc.add_paragraph(
    'The following third-party integrations and services have been identified through '
    'script analysis, network requests, and page inspection:'
)

integrations_data = [
    ['Integration Name', 'Type', 'Complexity', 'Purpose & Pages Used'],
    ['Adobe Experience Manager (AEM)', 'CMS Platform', 'High', 'Core CMS powering all pages. AEM clientlibs, components, and DAM assets throughout.'],
    ['Adobe Launch (DTM)', 'Tag Manager', 'Medium', 'Adobe Data Collection tag manager (assets.adobedtm.com). Manages all marketing tags.'],
    ['New Relic (Browser APM)', 'Monitoring', 'Low', 'Real User Monitoring (RUM) via nr-spa-1216.min.js. Performance tracking. All pages.'],
    ['FullStory', 'Session Analytics', 'Medium', 'Session replay (edge.fullstory.com). NPS surveys, site intercepts, datalayer. All pages.'],
    ['Google Tag Manager / GA4 / Ads', 'Analytics/Ads', 'Medium', 'GTM, GA4 (G-00W0WNR1CM), Google Ads (AW-964765464), DoubleClick (DC-4621208). All pages.'],
    ['Facebook Pixel', 'Advertising', 'Low', 'Facebook Events tracking (pixel: 662331570529793) with scroll depth and microdata. All pages.'],
    ['LinkedIn Insights', 'Advertising', 'Low', 'LinkedIn analytics tag (snap.licdn.com) for audience building. All pages.'],
    ['Coveo Search', 'Enterprise Search', 'High', 'AI search platform (platform-au.cloud.coveo.com). Powers site search and Help search.'],
    ['Auth0 (Authentication)', 'Identity/Auth', 'High', 'Auth0 SPA auth (welcome.auspost.com.au). MyPost login, session management. All pages.'],
    ['Google Maps Platform', 'Maps/Geo', 'High', 'Maps JS API + Places Autocomplete. Location finder on /locate and Parcel Locker pages.'],
    ['YouTube (Embedded)', 'Video', 'Low', 'YouTube iframe embeds for video content. Community/story pages.'],
    ['Vudoo', 'Interactive Media', 'Medium', 'Interactive content platform (vudoo.io). Enhanced media on community pages.'],
    ['Custom Analytics (dd.auspost)', 'Analytics', 'Medium', 'Custom analytics endpoint (dd.auspost.com.au/tags.js). Internal analytics aggregation.'],
    ['Covermore (Travel Insurance)', 'Partner API', 'High', 'Travel insurance quotes (auspost.poweredbycovermore.com). Partner integration.'],
    ['Western Union', 'Partner Service', 'Medium', 'International money transfer service integration. Money-travel section.'],
    ['Post BillPay', 'External Platform', 'Medium', 'Bill payment platform (www.postbillpay.com.au). External redirect.'],
    ['Salesforce (Help Portal)', 'CRM/Support', 'High', 'Salesforce help portal (helpandsupport.auspost.com.au). Forms, chatbot, case mgmt.'],
    ['DataGuard (Bot Protection)', 'Security', 'Medium', 'CAPTCHA/bot protection (geo.captcha-delivery.com). Conditional trigger.'],
    ['MyPost Tracking API', 'Internal API', 'High', 'Parcel tracking service (auspost.com.au/mypost/track). Real-time tracking. Homepage.'],
    ['Postage Calculator API', 'Internal API', 'High', 'Rate calculation engine. Dynamic pricing based on weight/size/destination.'],
]

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Medium Shading 1 Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table4.rows[0].cells
for i, header in enumerate(integrations_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in integrations_data[1:]:
    row_cells = table4.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_page_break()

# ============================================
# SECTION 5: COMPLEX USE CASES
# ============================================
doc.add_heading('5. Complex Use Cases & Observations', level=1)
doc.add_paragraph(
    'The following complex behaviours, edge cases, and functionality require special '
    'attention during migration:'
)

complex_data = [
    ['Use Case', 'Instances', 'Location', 'Why It Is Complex'],
    ['Authenticated User State (Auth0 SSO)', 'All pages (~377+)', 'Global header login/user menu', 'Auth0 silent auth runs on every page via iframe. User state changes nav options and personalized content. Requires SSO architecture decision.'],
    ['Coveo AI-Powered Search', '2 implementations', 'Header search + /help-and-support', 'Enterprise search with Headless engine, organization endpoints, platform URL config. Requires Coveo SDK and API keys.'],
    ['Multi-Level Mega Navigation', '1 (global)', 'All pages - header', '4-level nested nav with 3 audience segments, mobile responsive, keyboard accessible, context-sensitive help links per page.'],
    ['Google Maps Location Finder', '2-3 pages', '/locate, /parcel-lockers', 'Full Maps SDK + Places Autocomplete + geolocation + custom markers + location API. March 2025 pricing changes add cost.'],
    ['Real-Time Tracking Widget', '2+ pages', 'Homepage, /mypost/track/', 'Live parcel tracking with API calls. Authenticated and unauthenticated flows. Status polling.'],
    ['Postage Calculator (Dynamic Pricing)', '1 SPA', '/parcels-mail/calculate-postage-delivery-times/', 'Multi-step form, real-time pricing API, weight/size validators, service comparison, delivery date estimation. Heavy JS app.'],
    ['Currency Converter (Live Rates)', '1 page', '/currency-converter', 'Live exchange rate API with 40+ currencies. Searchable dropdown, real-time calculation.'],
    ['Live Chat / Chatbot', '1+ pages', '/help-and-support', 'AI chatbot (v1.14.5) with live agent escalation. Salesforce integration. Real-time messaging.'],
    ['Bot Protection / CAPTCHA', 'Conditional', 'All pages (triggered)', 'DataGuard CAPTCHA blocks automated access. Will interfere with migration scraping. Needs IP whitelisting.'],
    ['FullStory NPS Surveys', 'All pages (conditional)', 'Dynamic overlay/modal', 'Conditional NPS display based on cookies and probability. Config + usecase logic + view rendering.'],
    ['Context-Sensitive Help Links', 'All pages', 'Footer help link', 'Each page sets different help category code (DelNet100, DelLoc600, etc.) for contextual FAQ routing.'],
    ['Dynamic FAQ Loading', '~10+ pages', 'Service pages (FAQ sections)', 'FAQ content loaded dynamically (shows "Loading..." initially). Fetched from external CMS/Salesforce KB.'],
    ['Location-Specific Pages', '4000+ locations', '/locate/post-office/{state}/{suburb}/{postcode}', 'Individual location pages from database. Requires data feed and dynamic page generation system.'],
    ['Partner Service Redirects', '5+ integrations', 'BillPay, Covermore, Western Union', 'External partner platforms with session continuity and return-URL handling.'],
]

table5 = doc.add_table(rows=1, cols=4)
table5.style = 'Medium Shading 1 Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table5.rows[0].cells
for i, header in enumerate(complex_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in complex_data[1:]:
    row_cells = table5.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_page_break()

# ============================================
# SECTION 6: MIGRATION ESTIMATES
# ============================================
doc.add_heading('6. Migration Estimates', level=1)
doc.add_paragraph(
    'The following estimates are based on the identified templates, components, '
    'integrations, and complexity levels. Estimates assume a team of 2-3 frontend '
    'developers, 1 AEM/EDS specialist, and 1 QA engineer.'
)

doc.add_heading('6.1 Effort Breakdown', level=2)

effort_data = [
    ['Phase', 'Tasks', 'Effort (Days)', 'Notes'],
    ['Discovery & Planning', 'Template mapping, content model design, integration architecture, migration strategy', '10-15', 'Includes stakeholder workshops and technical design'],
    ['Infrastructure Setup', 'EDS project setup, block library creation, build pipeline, preview environment', '5-8', 'Foundation for all subsequent work'],
    ['Global Components Development', 'Header/mega-nav, footer, breadcrumbs, search integration (Coveo), Auth0 integration', '20-30', 'Highest complexity due to mega-nav and search/auth'],
    ['Block Development (Low Complexity)', 'Hero banner, CTA banner, quick links, feature lists, disclaimers, alert banners, breadcrumbs, app download', '10-12', '~8 blocks x 1-1.5 days each'],
    ['Block Development (Medium Complexity)', 'Cards, carousels, data tables, step-by-step, FAQ accordion, video embed, help cards, contact channels, shop promo', '18-25', '~9 blocks x 2-3 days each'],
    ['Block Development (High Complexity)', 'Tracking widget, location finder (Maps), currency converter, search (Coveo), suburb directory', '25-35', '~5 blocks x 5-7 days each (heavy API work)'],
    ['Template Assembly', 'Build all 11 template types using developed blocks, test responsive layouts', '15-20', 'Assembly + testing of template combinations'],
    ['Automated Content Migration', 'Script-based migration of ~200 standardized pages (policies, guides, directory, content)', '8-12', 'Import scripts + content validation + fixes'],
    ['Manual Content Migration', 'Hand-migration of ~177 pages with custom layouts, tools, integrations', '30-40', 'Custom page assembly and integration testing'],
    ['Integration Development', 'Auth0, Coveo, Google Maps, Tracking API, Calculator API, chat, analytics tags', '20-30', 'Complex API integrations requiring custom dev'],
    ['QA & Testing', 'Cross-browser, accessibility audit, performance, UAT, regression testing', '20-25', 'All pages across all breakpoints + integrations'],
    ['Content Freeze & Cutover', 'Final content sync, DNS cutover, go-live support, monitoring', '5-7', 'Coordinated cutover with business teams'],
    ['TOTAL ESTIMATED EFFORT', '', '186-259 days', 'Range accounts for unknowns and complexity variations'],
]

table6 = doc.add_table(rows=1, cols=4)
table6.style = 'Medium Shading 1 Accent 1'
table6.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table6.rows[0].cells
for i, header in enumerate(effort_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in effort_data[1:]:
    row_cells = table6.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()

doc.add_heading('6.2 Timeline Estimate', level=2)

timeline_data = [
    ['Phase', 'Duration', 'Team Size', 'Dependencies'],
    ['Discovery & Planning', 'Weeks 1-3', '3-4 people', 'Stakeholder availability, system access'],
    ['Infrastructure & Global Components', 'Weeks 3-8', '3 developers', 'Design tokens, brand assets, API credentials'],
    ['Block Development', 'Weeks 6-14', '2-3 developers', 'Approved designs, content models finalized'],
    ['Template Assembly & Content Migration', 'Weeks 12-20', '3-4 people', 'Blocks complete, content frozen'],
    ['Integration Development', 'Weeks 8-18', '2 developers', 'API access, partner coordination'],
    ['QA & Testing', 'Weeks 16-22', '2 QA engineers', 'Feature complete, test environments'],
    ['UAT & Go-Live', 'Weeks 22-24', 'Full team', 'Business sign-off, DNS control'],
    ['TOTAL TIMELINE', '~24 weeks (6 months)', '4-5 FTE average', 'Parallel workstreams assumed'],
]

table7 = doc.add_table(rows=1, cols=4)
table7.style = 'Medium Shading 1 Accent 1'
table7.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table7.rows[0].cells
for i, header in enumerate(timeline_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in timeline_data[1:]:
    row_cells = table7.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()

doc.add_heading('6.3 Cost Estimate Summary', level=2)

cost_data = [
    ['Category', 'Effort Range', 'Estimated Cost (AUD)*'],
    ['Automated Migration (scripts, tools, batch processing)', '25-35 days', '$50,000 - $70,000'],
    ['Manual / Custom Migration (pages, integrations, tools)', '95-135 days', '$190,000 - $270,000'],
    ['Block & Template Development', '68-100 days', '$136,000 - $200,000'],
    ['QA & Testing', '20-25 days', '$40,000 - $50,000'],
    ['Project Management & Coordination', '24 weeks x 0.5 FTE', '$48,000 - $60,000'],
    ['TOTAL ESTIMATED COST', '', '$464,000 - $650,000'],
]

table8 = doc.add_table(rows=1, cols=3)
table8.style = 'Medium Shading 1 Accent 1'
table8.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table8.rows[0].cells
for i, header in enumerate(cost_data[0]):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_data in cost_data[1:]:
    row_cells = table8.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

doc.add_paragraph()
doc.add_paragraph(
    '* Cost estimates assume blended rate of AUD $2,000/day per resource. '
    'Actual costs will vary based on team composition, location, and vendor rates.'
)

doc.add_heading('6.4 Key Risks & Assumptions', level=2)
risks = [
    'Bot protection (DataGuard CAPTCHA) may block automated content scraping - requires IP whitelisting or API access from Australia Post',
    'Auth0 integration complexity depends on whether existing tenant can be reused or new configuration is needed',
    'Coveo search requires licensed access and may need contract renewal/modification for new platform',
    'Google Maps API costs may increase with new implementation due to Places API pricing changes (March 2025)',
    '4000+ location pages require data feed access - cannot be manually migrated',
    'External partner integrations (Covermore, Western Union, BillPay) require coordination with third parties',
    'Content freeze period needed for migration cutover - business impact must be planned',
    'Dynamic FAQ content source (likely Salesforce Knowledge) must be maintained or migrated separately',
    'FullStory and analytics implementations require re-instrumentation in new architecture',
    'Mobile app deep links (auspost.app.link) routing must be maintained through migration',
]

for risk in risks:
    doc.add_paragraph(risk, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
doc.add_heading('Document End', level=2)
doc.add_paragraph(
    'This analysis was conducted in June 2026 based on the publicly accessible '
    'auspost.com.au website. Dynamic content, authenticated areas, and internal '
    'systems were not assessed. Estimates should be validated with Australia Post '
    'technical teams who have access to the full AEM authoring environment.'
)

# Save the document
output_path = '/backups/checksoundar/KyowaKirin-site-analysis/repo/AusPost_Site_Analysis_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
