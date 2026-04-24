#!/usr/bin/env python3
"""Generate comprehensive site analysis Word document for Krungthai Bank (krungthai.com)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS_DIR = "/tmp/playwright/screenshots"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_style(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '0072BC')
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if r % 2 == 0:
                set_cell_shading(cell, 'F0F8FF')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_screenshot(doc, filename, caption, width=5.5):
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ===================== COVER PAGE =====================
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Krungthai Bank Website Analysis')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 114, 188)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Comprehensive Site Analysis & Migration Assessment')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Target Site: https://krungthai.com/\nDate: April 2026\nPrepared for: AEM Edge Delivery Services Migration')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===================== TABLE OF CONTENTS =====================
    doc.add_heading('Table of Contents', level=1)
    for item in ['1. Executive Summary', '2. Templates Inventory', '3. Blocks / Components Catalog',
                  '4. Page Counts by Template', '5. Integrations Analysis',
                  '6. Complex Use Cases & Observations', '7. Migration Estimates', '8. Appendix: Screenshots']:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # ===================== 1. EXECUTIVE SUMMARY =====================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive analysis of the Krungthai Bank website (krungthai.com) '
        'for migration to Adobe Experience Manager (AEM) Edge Delivery Services. Krungthai Bank is one of '
        'Thailand\'s largest state-owned commercial banks, and its website serves multiple customer segments '
        '(Personal, SME, Corporate) with bilingual content (Thai/English), banking products, investor relations, '
        'sustainability reporting, and various financial tools.'
    )
    doc.add_paragraph()
    doc.add_heading('Key Findings', level=2)
    for f in [
        '14 distinct page templates identified across personal, corporate, IR, and sustainability sections',
        '22 reusable blocks/components cataloged with complexity ratings',
        'Bilingual site (Thai/English) with mirrored /th/ and /en/ URL structure',
        'Multiple customer segments with shared global navigation but segment-specific content',
        'Live exchange rate feed from external subdomain (exchangerate.krungthai.com)',
        'Loan calculator and financial planning tools requiring JS reimplementation',
        'Stock price widget via embedded iframe on Investor Relations pages',
        'Branch/ATM locator with province/district dropdown filtering and geolocation',
        'Heavy use of jQuery-based carousels (slick.js) across all page types',
        'Multiple external platform integrations (Internet Banking, Corporate Online, Krungthai BUSINESS, Careers portal, Complaint system)',
        'PDPA (Thailand Personal Data Protection Act) cookie consent banner on all pages',
        'Estimated ~350-500 content pages across Thai and English versions'
    ]:
        doc.add_paragraph(f, style='List Bullet')
    doc.add_page_break()

    # ===================== 2. TEMPLATES INVENTORY =====================
    doc.add_heading('2. Templates Inventory', level=1)
    doc.add_paragraph(
        'The following table lists all unique page templates identified across the Krungthai Bank website. '
        'The site uses a consistent layout framework with segment-specific variations.'
    )

    templates = [
        ['T01 - Personal Banking Homepage', 'High',
         'Hero carousel (5+ slides with auto-play), "Quick Help" service shortcuts carousel, promotions carousel, live exchange rate table, calculator tools section, financial partner articles carousel, news/updates carousel, social media links. Heavy JS.',
         'https://krungthai.com/th/personal'],
        ['T02 - Product Category Listing', 'Medium',
         'Hero carousel showcasing products, product card grid/carousel with interest rate and term details, loan calculator form, benefits USP icons, CTA to branch/app, related articles carousel, FAQ accordion (10+ items), breadcrumbs, social share.',
         'https://krungthai.com/th/personal/loan/personal-loan\nhttps://krungthai.com/th/personal/cards/debit-card'],
        ['T03 - Product Detail Page', 'High',
         'Product hero banner, tabbed content (features, eligibility, documents, fees), interest rate tables, step-by-step application guide, document download links, related products carousel, FAQ section, CTA buttons (Apply via app / at branch).',
         'https://krungthai.com/th/personal/loan/personal-loan/22\nhttps://krungthai.com/th/personal/deposits/211/640'],
        ['T04 - Corporate Banking Hub', 'Medium',
         'Same global layout as personal but with corporate-specific hero carousel, service categories (Cash Management, International, Business Loans), exchange rate widget, corporate news feed. Links to Krungthai BUSINESS platform.',
         'https://krungthai.com/th/corporate'],
        ['T05 - SME Landing Page', 'Medium',
         'SME-specific hero, product cards for SME loans and services, success stories carousel, financial education content.',
         'https://krungthai.com/th/content/sme'],
        ['T06 - Investor Relations Hub', 'High',
         'Hero banner, embedded stock price iframe widget, investor events calendar with date formatting, SET filing documents with PDF/ZIP downloads, Investor Kits tabbed by quarter (Q1-Q4), investor information cards carousel, chairman message section, contact details grid, FAQ accordion.',
         'https://krungthai.com/th/investor-relations'],
        ['T07 - News/Update Listing', 'Medium',
         'News cards with thumbnail, date badge, category tag (News/Announcement), pagination or carousel navigation, category filter tabs.',
         'https://krungthai.com/th/krungthai-update/news'],
        ['T08 - News/Article Detail', 'Medium',
         'Back navigation, category badge, article title (H1), publish date, hero image, rich text body with bullet lists and bold formatting, CTA buttons, related news carousel (5+ items), breadcrumbs, social share bar (Facebook, Line, Twitter, Embed).',
         'https://krungthai.com/th/krungthai-update/news-detail/3457'],
        ['T09 - Sustainability Hub', 'Medium',
         'CSR project hero carousel, sustainability pillars navigation cards carousel, CSR project gallery carousel, CSR news feed carousel, sustainability report download section with PDF link.',
         'https://krungthai.com/th/sustainability'],
        ['T10 - About Bank Hub', 'Medium',
         'Hero with vision/mission, quick-access cards (History, Board, Achievements, Structure), subsidiaries logo grid, corporate governance links, risk management section, career section.',
         'https://krungthai.com/th/about-ktb'],
        ['T11 - Content/Static Page', 'Low',
         'Simple rich text content page with header, body text, optional images, breadcrumbs, social share. Used for privacy policy, disclaimer, PDPA forms, legal pages.',
         'https://krungthai.com/th/content/privacy-policy\nhttps://krungthai.com/th/content/contact-us/disclaimer'],
        ['T12 - Branch/ATM Locator', 'High',
         'Province/district dropdown filters, business hours filter, geolocation "Find Nearby" button, results list by service type (Branch, ATM, FX Counter), tab for domestic vs. international branches. Map integration.',
         'https://krungthai.com/th/contact-us/ktb-location'],
        ['T13 - Rates & Fees Page', 'Medium',
         'Tabbed interest rate tables, fee schedules, exchange rate links, downloadable PDF documents. Data-heavy with structured tables.',
         'https://krungthai.com/th/rates'],
        ['T14 - Financial Calculator/Tool', 'High',
         'Interactive form with dropdowns (occupation, income, amount, term), real-time calculation engine, results display, financial planning tool with multiple scenarios.',
         'https://krungthai.com/th/financial-partner/calculator-tool/2\nhttps://krungthai.com/th/financial-partner/calculator/financial-planning'],
    ]

    add_table_with_style(doc, ['Template ID & Name', 'Complexity', 'Description', 'Reference URL(s)'], templates, col_widths=[4.5, 2, 7, 5])

    doc.add_paragraph()
    add_screenshot(doc, '01-homepage-full.png', 'Figure 1: Personal Banking Homepage (T01)', 3.5)
    doc.add_page_break()
    add_screenshot(doc, '02-product-listing.png', 'Figure 2: Product Category Listing - Personal Loans (T02)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '03-investor-relations.png', 'Figure 3: Investor Relations Hub (T06)', 3.5)
    doc.add_page_break()
    add_screenshot(doc, '04-news-article.png', 'Figure 4: News Article Detail (T08)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '05-sustainability.png', 'Figure 5: Sustainability Hub (T09)', 3.5)
    doc.add_page_break()

    # ===================== 3. BLOCKS / COMPONENTS CATALOG =====================
    doc.add_heading('3. Blocks / Components Catalog', level=1)
    doc.add_paragraph(
        'The following catalog identifies all reusable blocks and components across the Krungthai Bank site. '
        'Design variations of the same content model are noted as variants rather than separate blocks.'
    )

    blocks = [
        ['B01 - Global Header', 'High', 'Bank logo, Krungthai NEXT app download CTA, hamburger menu, mega-menu with 9 customer segments (Personal, SME, Corporate, Financial Partner, IR, Sustainability, About, Contact, Rates), language toggle (TH/EN), text size controls, search icon. Mobile: full-screen overlay menu with multi-level navigation.', 'All pages'],
        ['B02 - Global Footer', 'Medium', 'DPA (Data Protection) badge, social media icons (Facebook, Line, X/Twitter, Instagram, YouTube, TikTok), collapsible accordion sections (Sustainability, Careers, Krungthai Update, PDPA, Others, Contact, Disclosure), copyright with SWIFT code.', 'All pages'],
        ['B03 - Hero Carousel', 'High', 'Full-width image carousel with auto-play, prev/next arrows, dot pagination. Responsive images (desktop/mobile variants). Overlay text with title, description, and CTA link. Used on homepage, product pages, sustainability, IR. Multiple design variants.', 'Homepage, Product, Sustainability, IR pages'],
        ['B04 - Product Card Carousel', 'High', 'Horizontal scrolling product cards. Each card: product image, title, description, interest rate display, loan term, "Read Details" CTA. Dot pagination. Used for loans, deposits, cards, insurance products.', 'Product listing pages'],
        ['B05 - Promotion Card Carousel', 'Medium', 'Image thumbnail cards with title overlay. Links to promotion detail pages. Horizontal scroll with dot navigation. "View All" link.', 'Homepage, product pages'],
        ['B06 - Exchange Rate Table', 'High', 'Live currency rate display with flag icons (USD, GBP, EUR, JPY, HKD), buy/sell columns, last update timestamp, "View All" link to external exchangerate.krungthai.com. Real-time data feed.', 'Homepage (Personal, Corporate)'],
        ['B07 - Loan Calculator Form', 'High', 'Interactive form: occupation dropdown, monthly income input, desired amount input, repayment period dropdown (12-60 months), "View Calculation" submit button. Client-side JS calculation. Results display area.', 'Personal loan listing page'],
        ['B08 - FAQ Accordion', 'Low', 'Expandable Q&A pairs (10+ items typical). Each item has question header with toggle button, collapsible answer panel. Schema.org FAQ markup for SEO. "View All FAQ" link.', 'Product pages, IR page'],
        ['B09 - Quick Help / Service Shortcuts', 'Medium', 'Circular icon cards carousel with service labels (Online Services, Fees, Debit Card, Transaction Code, Loans, etc.). "Services you may be interested in" heading. Dot pagination.', 'Homepage'],
        ['B10 - News Card Carousel', 'Medium', 'News item cards with thumbnail image, date badge with category (News/Announcement), headline text. Horizontal carousel with dot pagination. "Read Details" link.', 'Homepage, news listing, article detail (related)'],
        ['B11 - Financial Partner Articles', 'Medium', 'Article cards with thumbnail, category badge (SME Tips, Loans, etc.), article title. Horizontal carousel. "View All" link.', 'Homepage, product pages'],
        ['B12 - Breadcrumb Navigation', 'Low', 'Hierarchical path with ">" separators. Links to parent pages. Current page at end.', 'All interior pages'],
        ['B13 - Social Share Bar', 'Low', 'Facebook, Line, Twitter/X, Embed link sharing icons. "Share this page" toggle. Horizontal layout.', 'All content/detail pages'],
        ['B14 - Cookie Consent Banner', 'Low', 'PDPA-compliant bottom overlay. Accept all / Cookie settings buttons. Link to disclaimer. Privacy dialog with consent management.', 'All pages (first visit)'],
        ['B15 - Contact/CTA Bar', 'Low', 'Fixed bottom bar with "Find Branch/Phone" and "Contact Bank" icon links. Links to branch locator and contact page.', 'Product pages'],
        ['B16 - Investor Kits Tabs', 'Medium', 'Quarterly tab navigation (Q1-Q4). Each tab shows document thumbnails (Factsheet, Analyst Presentation, MD&A, Financial Statements) with download links. "Download All" button.', 'Investor Relations hub'],
        ['B17 - Stock Price Widget', 'High', 'Embedded iframe displaying live KTB stock price from external source. Real-time data feed.', 'Investor Relations hub'],
        ['B18 - Investor Calendar', 'Medium', 'Event list with formatted dates (day, month, year in Thai), event description. "View All" link to full calendar.', 'Investor Relations hub'],
        ['B19 - SET Filing Documents', 'Medium', 'Document cards with thumbnail, title, download icon. PDF and ZIP file types. Links to /Download/ directory.', 'Investor Relations hub'],
        ['B20 - Benefits/USP Icons Row', 'Low', 'Horizontal row of icon + text pairs highlighting product benefits (e.g., High Credit, Easy Apply, Low Interest, No Guarantor). Typically 4 items.', 'Product listing pages'],
        ['B21 - App Download CTA Block', 'Medium', 'Featured image of loan product, title, description, disclaimer text, "Application Steps" CTA button. Promotes Krungthai NEXT app for online applications.', 'Product listing pages'],
        ['B22 - Sustainability Report Download', 'Medium', 'Report cover image, title, year, download CTA button (PDF link), "View All" link to reports listing. Background image styling.', 'Sustainability hub'],
    ]

    add_table_with_style(doc, ['Block ID & Name', 'Complexity', 'Description & Functionality', 'Reference URL(s)'], blocks, col_widths=[4, 2, 8, 4.5])
    doc.add_page_break()

    # ===================== 4. PAGE COUNTS BY TEMPLATE =====================
    doc.add_heading('4. Page Counts by Template', level=1)
    doc.add_paragraph('Estimated page counts based on navigation structure, URL patterns, and content analysis. Counts include both Thai (/th/) and English (/en/) versions.')

    page_counts = [
        ['T01 - Personal Banking Homepage', '2 (TH+EN)', 'Manual', 'Complex multi-carousel layout, live data feeds, personalized recommendations.'],
        ['T02 - Product Category Listing', '~20-25', 'Semi-Automatic', 'Deposits, Cards, Loans, Insurance, Investment, Cash Mgmt categories. Standardized template with calculator variants.'],
        ['T03 - Product Detail Page', '~80-120', 'Semi-Automatic', 'Individual banking products (each loan type, card, deposit product, insurance plan). Structured data can be imported but tabbed content needs manual review.'],
        ['T04 - Corporate Banking Hub', '2 (TH+EN)', 'Manual', 'Corporate-specific layout with different product categories and CTA flows.'],
        ['T05 - SME Landing Page', '2 (TH+EN)', 'Manual', 'Custom layout for SME segment.'],
        ['T06 - Investor Relations Hub', '2 (TH+EN)', 'Manual', 'Complex: stock widget iframe, calendar, document downloads, tabbed kits.'],
        ['T07 - News/Update Listing', '~5-8', 'Semi-Automatic', 'News listing, promotion listing, announcement listing, CSR news listing.'],
        ['T08 - News/Article Detail', '~200-300+', 'Automatic', 'Standardized article content with image + rich text. Ideal for bulk import. IDs go to 3400+.'],
        ['T09 - Sustainability Hub + Sub-pages', '~15-20', 'Semi-Automatic', 'Hub + ESG policy pages, CSR projects (~20 items), growing together sub-pages.'],
        ['T10 - About Bank Hub + Sub-pages', '~15-20', 'Semi-Automatic', 'Hub + history, org chart, board members, governance, risk management, career pages.'],
        ['T11 - Content/Static Pages', '~20-30', 'Semi-Automatic', 'Privacy, disclaimer, PDPA forms, sitemap, whistle-blowing, legal. Simple rich text.'],
        ['T12 - Branch/ATM Locator', '2 (TH+EN)', 'Manual', 'Complex: Province/district filters, geolocation, map integration, results display.'],
        ['T13 - Rates & Fees Page', '~5-8', 'Manual', 'Data tables, interest rate schedules, downloadable PDFs. Data frequently updated.'],
        ['T14 - Financial Calculator/Tool', '~3-5', 'Manual', 'Custom JS calculation engines. Loan calculator, financial planning tool, FX tools.'],
    ]

    add_table_with_style(doc, ['Template', 'Est. Page Count', 'Migration Type', 'Notes / Rationale'], page_counts, col_widths=[4, 2.5, 3, 9])

    doc.add_paragraph()
    doc.add_heading('Summary Totals', level=2)
    summary = [
        ['Automatic Migration (bulk import)', '~200-300', 'News articles, promotion detail pages, CSR news articles'],
        ['Semi-Automatic (template + review)', '~150-200', 'Product pages, product listings, about sub-pages, sustainability sub-pages, static content'],
        ['Manual Migration (custom development)', '~30-40', 'Homepages, IR hub, branch locator, calculators, rates pages, corporate hub'],
        ['TOTAL ESTIMATED PAGES', '~380-540', 'All content across Thai and English versions'],
    ]
    add_table_with_style(doc, ['Migration Category', 'Est. Pages', 'Includes'], summary, col_widths=[5, 3, 10])
    doc.add_page_break()

    # ===================== 5. INTEGRATIONS ANALYSIS =====================
    doc.add_heading('5. Integrations Analysis', level=1)
    integrations = [
        ['jQuery / jQuery Migrate 1.4.1', 'Custom Code / Framework', 'Medium', 'Core JS framework powering all interactive elements. jQuery Migrate for backward compatibility. Must be replaced with vanilla JS or modern framework for EDS.', 'All pages'],
        ['Slick.js Carousel', 'Plugin / UI', 'Medium', 'Used for all carousels site-wide (hero, products, promotions, news, articles). Multiple instances per page with different configurations.', 'All pages with carousels'],
        ['Exchange Rate Feed', 'API / External', 'High', 'Live currency rates from exchangerate.krungthai.com. Displays USD, GBP, EUR, JPY, HKD with buy/sell prices and timestamp.', 'Homepage (Personal + Corporate)'],
        ['Stock Price Widget (iframe)', 'Embed / External', 'High', 'Embedded iframe showing real-time KTB stock price data. Source appears to be SET (Stock Exchange of Thailand) or third-party provider.', 'https://krungthai.com/th/investor-relations'],
        ['Krungthai NEXT App Integration', 'Embed / Mobile', 'Medium', 'QR code for app download, deep links to app features. Prominent CTA across site for online loan applications and banking.', 'All pages (header CTA)'],
        ['Facebook Share API', 'API / Social', 'Low', 'Share functionality via Facebook sharer URL. Social media link in footer.', 'All content pages, footer'],
        ['LINE Share API', 'API / Social', 'Low', 'Share functionality via lineit.line.me share URL. Footer link to LINE official account.', 'All content pages, footer'],
        ['Twitter/X Share API', 'API / Social', 'Low', 'Share functionality via twitter.com share URL.', 'All content pages'],
        ['PDPA Cookie Consent Manager', 'Custom Code / Compliance', 'Medium', 'Thailand Personal Data Protection Act compliant consent banner. Accept/Settings dialog. Required for legal compliance.', 'All pages'],
        ['Internet Banking (ktbnetbank.com)', 'External Platform', 'Low', 'External link to consumer Internet Banking portal. No embedded integration.', 'Footer, E-Banking menu items'],
        ['Corporate Online (newcb.ktb.co.th)', 'External Platform', 'Low', 'External link to corporate banking portal.', 'Corporate section, footer'],
        ['Krungthai BUSINESS (business.krungthai.com)', 'External Platform', 'Low', 'External link to business banking platform. Featured in corporate news articles.', 'Corporate section, news articles'],
        ['Careers Portal (careers.hr-krungthaigroup.com)', 'External Platform', 'Low', 'External HR/recruitment system. Link in footer and About section.', 'Footer, About Bank section'],
        ['Complaint System (complaint.krungthai.com)', 'External Platform', 'Low', 'External complaint/feedback submission system.', 'Contact section, footer'],
        ['E-Procurement (eprocurement.krungthai.com)', 'External Platform', 'Low', 'External procurement platform link.', 'Footer'],
        ['NACC ITA Assessment (itas.nacc.go.th)', 'External / Government', 'Low', 'Link to National Anti-Corruption Commission assessment.', 'Footer disclosure section'],
        ['PDF/ZIP Document Downloads', 'Custom Code / CDN', 'Medium', 'Document serving from /Download/ directory. Financial statements (ZIP), press releases (PDF), sustainability reports, investor kits.', 'IR, Sustainability, About pages'],
    ]
    add_table_with_style(doc, ['Integration Name', 'Type', 'Complexity', 'Description', 'Reference URL(s)'], integrations, col_widths=[3.5, 2.5, 2, 7, 3.5])
    doc.add_page_break()

    # ===================== 6. COMPLEX USE CASES =====================
    doc.add_heading('6. Complex Use Cases & Observations', level=1)
    complex_cases = [
        ['Bilingual Content (Thai/English)', 'Entire site (~380-540 pages x2)', 'All pages (/th/ and /en/ paths)',
         'Full site mirrored in Thai and English with /th/ and /en/ URL prefixes. Navigation, content, and product details must be maintained in both languages.',
         'Requires i18n content model. Consider EDS language copies or translation workflow integration.'],
        ['Live Exchange Rate Feed', '2 pages (Personal + Corporate homepages)', 'Homepage exchange rate table',
         'Real-time currency data from exchangerate.krungthai.com with buy/sell rates and timestamps. Updates throughout the day.',
         'Needs API integration or server-side data fetching. Cannot be static content. Consider client-side fetch or edge function.'],
        ['Stock Price Widget (iframe)', '1 page (IR hub)', 'https://krungthai.com/th/investor-relations',
         'Embedded iframe displaying real-time KTB stock data. External dependency on stock data provider.',
         'iframe embed approach works for EDS. Need to confirm cross-origin security and responsive sizing.'],
        ['Loan Calculator Engine', '~3 pages', 'Product listing pages (Personal Loans)',
         'Client-side calculation with occupation dropdown, income, amount, and term inputs. Real-time computation of monthly payments. Complex financial formulas.',
         'Must reimplemented as EDS block with vanilla JS. Calculation logic must be precisely replicated.'],
        ['Branch/ATM Locator', '2 pages (TH+EN)', 'https://krungthai.com/th/contact-us/ktb-location',
         'Multi-criteria search: 77 provinces, districts, business hours. Geolocation "Find Nearby" button. Results categorized by service type. Map display.',
         'Most complex page. Needs backend API for location data, possible Google Maps integration, geolocation API.'],
        ['Multi-segment Navigation', 'All pages', 'Global header',
         '9 customer segments (Personal, SME, Corporate, Financial Partner, IR, Sustainability, About, Contact, Rates) with deep multi-level mega-menu. Mobile: full-screen overlay with back-navigation. JavaScript-driven menu state.',
         'Complex mega-menu block needed. Segment switching changes available sub-navigation. Consider conditional content loading.'],
        ['Investor Kits Quarterly Tabs', '1 page', 'IR hub',
         'Tabbed document section with Q1-Q4 navigation. Each tab shows 4 document types (Factsheet, Analyst Presentation, MD&A, Financial Statements) with download links.',
         'EDS tabbed block with dynamic content per quarter. Document management for quarterly updates.'],
        ['PDPA Compliance Requirements', 'All pages', 'Cookie consent banner, privacy policy',
         'Thailand Personal Data Protection Act (PDPA) mandates consent management, privacy policy, data subject rights forms. Cookie banner with settings dialog.',
         'Legal requirement. Must implement PDPA-compliant consent management. DPA badge in footer.'],
        ['Carousel-Heavy Design Pattern', '~30+ pages', 'All major page types',
         'Every major page uses multiple slick.js carousels: hero, products, promotions, articles, news. Some pages have 5+ carousels. Performance concern.',
         'Replace with EDS-native lightweight carousel block. Consider reducing carousel count for performance. Critical for Core Web Vitals.'],
        ['Document Download System', '~10 pages', 'IR, Sustainability, About',
         'Financial statements (ZIP), press releases (PDF), sustainability reports, corporate governance documents served from /Download/ path. Version management by quarter/year.',
         'Need DAM integration for document management. Consider EDS document listing block with proper metadata.'],
    ]
    add_table_with_style(doc, ['Use Case', 'Instances / Scope', 'Where Found', 'Description', 'Why Complex / Migration Impact'], complex_cases, col_widths=[3.5, 2.5, 3, 5, 4.5])
    doc.add_page_break()

    # ===================== 7. MIGRATION ESTIMATES =====================
    doc.add_heading('7. Migration Estimates', level=1)
    doc.add_paragraph('Estimates assume a team of 2-3 EDS developers, 1 content migration specialist, and 1 QA engineer.')

    doc.add_heading('7.1 Effort Breakdown by Phase', level=2)
    effort = [
        ['Phase 1: Discovery & Design System', '', '', ''],
        ['  Design token extraction (KTB blue palette, typography)', '2-3 days', 'Semi-Automatic', 'Extract Krungthai brand colors, fonts, spacing from existing CSS'],
        ['  Template architecture design (14 templates)', '3-5 days', 'Manual', 'Map templates to EDS page structures with bilingual support'],
        ['  Block specification & design (22 blocks)', '4-5 days', 'Manual', 'Define block specs with content models for Thai/English'],
        ['', '', '', ''],
        ['Phase 2: Block Development', '', '', ''],
        ['  Global Header (mega-menu, 9 segments)', '5-6 days', 'Manual', 'Most complex block with multi-level navigation and segment switching'],
        ['  Global Footer (accordion sections)', '2 days', 'Manual', 'Collapsible sections with DPA badge'],
        ['  Hero Carousel (3+ variants)', '3-4 days', 'Manual', 'Auto-play, responsive images, overlay text, multiple layouts'],
        ['  Product Card Carousel', '3-4 days', 'Manual', 'Product data display with interest rates and terms'],
        ['  Exchange Rate Table', '3-4 days', 'Manual', 'API integration for live data, currency flags, formatting'],
        ['  Loan Calculator Form', '4-5 days', 'Manual', 'Complex financial calculation engine in vanilla JS'],
        ['  FAQ Accordion', '1 day', 'Manual', 'Schema.org FAQ markup support'],
        ['  Stock Price Widget', '1-2 days', 'Manual', 'iframe embed with responsive sizing'],
        ['  Investor Kits Tabs + Calendar + Documents', '3-4 days', 'Manual', 'Tabbed quarterly content, date formatting, file downloads'],
        ['  Branch Locator', '5-7 days', 'Manual', 'Province/district filters, geolocation, results display, map'],
        ['  Remaining blocks (10+)', '6-8 days', 'Manual', 'Social share, breadcrumbs, CTA bar, USP icons, app download, etc.'],
        ['', '', '', ''],
        ['Phase 3: Content Migration', '', '', ''],
        ['  News articles (~200-300)', '3-5 days', 'Automatic', 'Bulk import of standardized articles'],
        ['  Product pages (~80-120)', '5-7 days', 'Semi-Automatic', 'Product data mapping + tabbed content review'],
        ['  Product category pages (~20-25)', '3-4 days', 'Semi-Automatic', 'Template application + calculator/FAQ configuration'],
        ['  IR, Sustainability, About sub-pages (~50-60)', '4-5 days', 'Semi-Automatic', 'Document downloads, sub-page content migration'],
        ['  Static/content pages (~20-30)', '2-3 days', 'Semi-Automatic', 'Simple rich text content'],
        ['  English language versions', '5-7 days', 'Semi-Automatic', 'Mirror Thai content structure to English with translated content'],
        ['  Homepages + hub pages (~10)', '3-4 days', 'Manual', 'Complex multi-block assembly'],
        ['', '', '', ''],
        ['Phase 4: Integration Setup', '', '', ''],
        ['  Exchange rate API integration', '2-3 days', 'Manual', 'API endpoint mapping, data formatting'],
        ['  Stock price widget configuration', '1 day', 'Manual', 'iframe embed setup and testing'],
        ['  Social media sharing', '1 day', 'Manual', 'Facebook, LINE, Twitter share configuration'],
        ['  PDPA cookie consent', '2 days', 'Manual', 'Compliance consent management implementation'],
        ['  Document download system', '2-3 days', 'Manual', 'DAM setup for PDF/ZIP serving'],
        ['  Analytics setup', '1-2 days', 'Manual', 'Tag configuration'],
        ['', '', '', ''],
        ['Phase 5: QA & Testing', '', '', ''],
        ['  Template/block visual QA', '5-7 days', 'Manual', 'Cross-browser, responsive, bilingual testing'],
        ['  Content validation (Thai + English)', '3-5 days', 'Semi-Automatic', 'Bilingual content accuracy verification'],
        ['  Integration testing', '2-3 days', 'Manual', 'API endpoints, live data feeds, downloads'],
        ['  Performance testing', '2-3 days', 'Manual', 'Lighthouse, Core Web Vitals optimization'],
        ['  UAT support', '3-5 days', 'Manual', 'Stakeholder review and feedback cycles'],
    ]
    add_table_with_style(doc, ['Task / Activity', 'Estimated Duration', 'Migration Type', 'Notes'], effort, col_widths=[6, 3, 3, 6.5])

    doc.add_paragraph()
    doc.add_heading('7.2 Total Estimated Schedule', level=2)
    schedule = [
        ['Phase 1: Discovery & Design System', '2 weeks', '72-104 hrs'],
        ['Phase 2: Block Development', '5-7 weeks', '280-400 hrs'],
        ['Phase 3: Content Migration', '4-5 weeks', '200-280 hrs'],
        ['Phase 4: Integration Setup', '1.5-2 weeks', '72-96 hrs'],
        ['Phase 5: QA & Testing', '3-4 weeks', '120-184 hrs'],
        ['Buffer / Contingency (15%)', '2-3 weeks', '112-160 hrs'],
        ['TOTAL', '17-23 weeks (~4-5.5 months)', '856-1,224 hrs'],
    ]
    add_table_with_style(doc, ['Phase', 'Duration', 'Estimated Hours'], schedule, col_widths=[7, 5, 5])

    doc.add_paragraph()
    doc.add_heading('7.3 Recommended Migration Approach', level=2)
    approach = [
        ['Wave 1 (Weeks 1-6)', 'Foundation + Quick Wins',
         'Design system, global header/footer, news article bulk import (~200-300 pages Thai), static content pages. Delivers ~250 pages.'],
        ['Wave 2 (Weeks 5-12)', 'Banking Products',
         'Product card carousel, calculator, FAQ blocks. Product category + detail pages (~100-150). Personal and Corporate hubs. Delivers ~160 pages.'],
        ['Wave 3 (Weeks 10-16)', 'Investor Relations & Sustainability',
         'IR hub with stock widget, kits tabs, documents. Sustainability hub with CSR content. About Bank pages. Delivers ~50 pages.'],
        ['Wave 4 (Weeks 14-18)', 'Complex Features + English',
         'Branch locator, exchange rate feed, English language version mirror. Delivers ~200+ pages (EN mirror).'],
        ['Wave 5 (Weeks 16-21)', 'QA, UAT & Launch',
         'Comprehensive bilingual testing, performance optimization, stakeholder review, go-live preparation.'],
    ]
    add_table_with_style(doc, ['Wave / Timeline', 'Focus Area', 'Scope & Deliverables'], approach, col_widths=[3.5, 3.5, 11.5])

    doc.add_page_break()

    # ===================== 8. APPENDIX =====================
    doc.add_heading('8. Appendix: Additional Screenshots', level=1)
    add_screenshot(doc, '01-homepage.png', 'Figure A1: Homepage - Above the Fold (Viewport)', 5)

    # Save
    output_path = '/workspace/Krungthai_Bank_Site_Analysis_Report.docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')

if __name__ == '__main__':
    main()
