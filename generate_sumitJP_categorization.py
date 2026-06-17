#!/usr/bin/env python3
"""Generate the sumitclub.jp URL categorization & automated-migration-feasibility
report (Word document) from the live DOM analysis of the provided URL list.

Data source: /workspace/current/.migration/sumitJP/report-data.json (per-category
counts + migration mode), confirmed against live DOM inspection (dom-inspection.json).
"""
import json
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA = '/workspace/current/.migration/sumitJP/report-data.json'
OUT = '/workspace/current/SumitclubJP_URL_Categorization_Migration_Analysis.docx'
TABLE_STYLE = 'Medium Shading 1 Accent 1'

rows = json.load(open(DATA))

# DOM-template family grouping (confirmed via live inspection).
FAMILY = {
    'Standard AEM Content Shell (/ja/)': [
        'Usage / How-to Page', 'Travel Service Page', 'Insurance Product/Info',
        'Point Program Page', 'Legal / Policy Text', 'Contact / Support',
        'Campaign Landing Page', 'Entertainment/Lifestyle Page', 'Gourmet Service Page',
        'Club Online Promo (unique)', 'Section Landing Page', 'Section Index',
        'Commercial Card Page', 'Other Content Page', 'Notice / News Detail',
        'Notice / News Index', 'Sitemap',
    ],
    'AEM Card-Detail / Listing Shell': [
        'Card Product Detail', 'Card Lineup Listing/Index',
    ],
    'AEM Corporate Shell (/ja/corporate_site)': [
        'Corporate Info (AEM)', 'Corporate News List', 'Corporate Special/Anniversary',
    ],
    'Legacy Corporate Microsite (/corporate)': [
        'Corporate Info (legacy shell)', 'Corporate Kaizen Notice (archive)',
        'Corporate Recruit Page',
    ],
    'Legacy English Shell (/en)': [
        'Homepage (locale)', 'Service Page (legacy en)', 'Announcement Page',
        'FAQ Page', 'Info Index', 'HTTP Error Page',
    ],
    'Standalone Microsites (non-AEM)': [
        'Card Application Landing Page (LP)', 'Identity Verification Microsite',
        'Club Online Bumper Page', 'Cancellation/Termination Flow',
    ],
    'Auth-gated / System (excluded)': [
        'Member Dining-Selection Detail (R####)', 'Member Dining-Selection Index/Search',
        'Login/Member System Page', 'Sign-on / Login System', 'System Error Stub',
    ],
    'Non-Page Assets (excluded)': [
        'SSI Include Fragment', 'LP Partial Fragment', 'Bumper/Redirect Stub',
        'Blank/Helper Stub', 'Tracking/System Stub', 'Modal/Popup Fragment',
        'Auth Widget Fragment', 'Redirect / Router Stub', 'App-bridge Stub',
    ],
    'E-statement Email Templates (separate track)': [
        'E-statement Email Template',
    ],
}
cat_to_family = {c: f for f, cs in FAMILY.items() for c in cs}

# Build doc -----------------------------------------------------------------
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def add_table(data):
    t = doc.add_table(rows=1, cols=len(data[0]))
    t.style = TABLE_STYLE
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(data[0]):
        t.rows[0].cells[i].text = str(h)
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in data[1:]:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


total = sum(r['n'] for r in rows)
R_COUNT = 1004  # member dining-selection R####.html detail pages in the provided list

# Title
doc.add_paragraph(); doc.add_paragraph()
ti = doc.add_heading('TRUST CLUB (www.sumitclub.jp)', level=0)
ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
st = doc.add_heading('URL Categorization & Automated Migration Feasibility', level=1)
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: June 2026\n').bold = True
info.add_run('Source list: sumitJP-URLs.txt\n')
info.add_run(f'Distinct URLs analysed: {total} (excludes ~{R_COUNT} member R#### detail pages, counted separately)\n')
info.add_run('Method: path-pattern classification confirmed by live DOM inspection\n')
info.add_run('Platform: Adobe Experience Manager (AEM) + several standalone microsites')
doc.add_page_break()

# TOC
doc.add_heading('Table of Contents', level=1)
for s in ['1. Approach & Method', '2. Migration-Mode Summary',
          '3. DOM Template Families', '4. Full Category Breakdown',
          '5. Excluded (Non-Page) Assets', '6. Notes & Recommendations']:
    doc.add_paragraph(s, style='List Number')
doc.add_page_break()

# 1. Approach
doc.add_heading('1. Approach & Method', level=1)
doc.add_paragraph(
    'Every URL in the supplied list was classified by its path pattern into a content '
    'category, then a representative page from each category was opened in a real browser '
    'to confirm its DOM structure (heading hierarchy, sections, blocks such as carousels, '
    'tables, accordions, tabs and card grids). Categories that share the same underlying '
    'DOM template were then grouped into "template families". Each category carries a '
    'migration-mode tag:')
for lab, desc in [
    ('Automated', 'Standardized DOM; can be bulk-imported by a block parser + page transformer with spot-check QA.'),
    ('Assisted', 'Same shell but per-page layout variation; bulk-imported then reviewed/cleaned per page.'),
    ('Manual', 'Custom layout, marketing-critical, or microsite outside the main AEM shell; rebuilt by hand.'),
    ('Exclude (not a page)', 'SSI include fragments, partials, blank/helper, tracking and redirect stubs — not standalone pages.'),
    ('Exclude / Re-point', 'Login/sign-on, router and system stubs — re-pointed to their external systems, not migrated as content.'),
    ('Exclude / Separate track', 'E-statement email templates — belong to the email programme, not the website migration.'),
    ('Manual / Auth-gated', 'Behind member login; not anonymously reachable, handled as a single repeating template.'),
    ('Manual / Replace', 'HTTP error pages — replaced by native EDS error handling.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{lab}: ').bold = True
    p.add_run(desc)
doc.add_page_break()

# 2. Migration-mode summary
doc.add_heading('2. Migration-Mode Summary', level=1)
mode_tot = {}
for r in rows:
    mode_tot[r['mode']] = mode_tot.get(r['mode'], 0) + r['n']
mode_rows = [['Migration Mode', 'Distinct URLs', '% of analysed']]
for m, n in sorted(mode_tot.items(), key=lambda x: -x[1]):
    mode_rows.append([m, n, f'{100*n/total:.1f}%'])
mode_rows.append(['TOTAL (distinct)', total, '100%'])
add_table(mode_rows)
doc.add_paragraph()
migratable = sum(n for m, n in mode_tot.items() if m.startswith(('Automated', 'Assisted', 'Manual')) and 'Auth' not in m)
excluded = total - migratable
doc.add_paragraph(
    f'Of {total} distinct URLs, roughly {migratable} are genuine migratable website pages; '
    f'about {excluded} are non-page assets, auth-gated/system pages, or email templates that are '
    'excluded from the content migration or handled on a separate track. In addition, the list '
    f'contains ~{R_COUNT} member "dining-selection" detail pages (…/search/R####.html) which are '
    'behind member login and JS-driven — these are treated as ONE repeating, auth-gated template, '
    'not 1,004 distinct builds.')
doc.add_page_break()

# 3. Families
doc.add_heading('3. DOM Template Families', level=1)
doc.add_paragraph(
    'Categories grouped by shared DOM template. The single biggest opportunity is the standard '
    '/ja/ AEM content shell, which underlies most usage, travel, insurance, point, legal, notice '
    'and landing pages — one block set and one importer cover them all.')
fam_tot = {}
for r in rows:
    f = cat_to_family.get(r['cat'], 'Other')
    fam_tot[f] = fam_tot.get(f, 0) + r['n']
fam_rows = [['Template Family', 'Distinct URLs', 'Predominant Mode']]
fam_mode = {
    'Standard AEM Content Shell (/ja/)': 'Automated / Assisted',
    'AEM Card-Detail / Listing Shell': 'Manual / Assisted',
    'AEM Corporate Shell (/ja/corporate_site)': 'Assisted',
    'Legacy Corporate Microsite (/corporate)': 'Assisted (kaizen: Automated)',
    'Legacy English Shell (/en)': 'Assisted / Manual',
    'Standalone Microsites (non-AEM)': 'Manual',
    'Auth-gated / System (excluded)': 'Exclude / Re-point',
    'Non-Page Assets (excluded)': 'Exclude (not a page)',
    'E-statement Email Templates (separate track)': 'Separate track',
}
for f, n in sorted(fam_tot.items(), key=lambda x: -x[1]):
    fam_rows.append([f, n, fam_mode.get(f, '—')])
add_table(fam_rows)
doc.add_page_break()

# 4. Full category breakdown
doc.add_heading('4. Full Category Breakdown', level=1)
doc.add_paragraph('All 43 categories with counts, migration mode, template family and an example URL.')
cat_rows = [['Category', 'Count', 'Mode', 'Family', 'Example URL']]
for r in sorted(rows, key=lambda x: -x['n']):
    cat_rows.append([r['cat'], r['n'], r['mode'],
                     cat_to_family.get(r['cat'], 'Other'),
                     r['ex'].replace('https://www.sumitclub.jp', '')])
add_table(cat_rows)
doc.add_page_break()

# 5. Excluded assets
doc.add_heading('5. Excluded (Non-Page) Assets', level=1)
doc.add_paragraph(
    'These URLs are not standalone pages and were live-confirmed as bare partials, stubs or '
    'system endpoints. They should NOT be counted as migration units. SSI fragments (e.g. '
    '/ja/ssi/footer.html) have an empty <head> and only a markup snippet; e-statement emails '
    'carry noindex/nofollow and no site chrome.')
exc_rows = [['Category', 'Count', 'Why excluded']]
why = {
    'E-statement Email Template': 'Email markup (noindex), no site chrome — email programme, separate track.',
    'SSI Include Fragment': 'Server-side include snippet (header/footer/side-nav) — becomes EDS block, not a page.',
    'LP Partial Fragment': 'Landing-page partials (_offer/_header/parts) assembled into a parent LP.',
    'Redirect / Router Stub': 'rd_/redirect/WD router stubs — re-pointed, not migrated.',
    'Bumper/Redirect Stub': 'Interstitial bumper link pages.',
    'Modal/Popup Fragment': 'Modal/popup markup loaded into a parent page.',
    'Blank/Helper Stub': 'Blank iframe / javascript-off helper pages.',
    'Sign-on / Login System': 'Authentication entry points — external system.',
    'Tracking/System Stub': 'Tracking pixels / verification / log endpoints.',
    'System Error Stub': 'Point-mall system-failure stubs.',
    'App-bridge Stub': 'Native-app bridge endpoint.',
}
for r in sorted(rows, key=lambda x: -x['n']):
    if r['mode'].startswith('Exclude'):
        exc_rows.append([r['cat'], r['n'], why.get(r['cat'], 'Non-page asset.')])
add_table(exc_rows)
doc.add_page_break()

# 6. Notes
doc.add_heading('6. Notes & Recommendations', level=1)
for n in [
    'Consolidate the standard /ja/ AEM content shell into a single EDS template + shared block '
    'set; it covers the large majority of genuine pages (usage, travel, insurance, point, legal, '
    'notice, campaign, landing) and is the primary automation lever.',
    'The ~1,004 member dining-selection detail pages (…/search/R####.html) are behind member '
    'login and rendered via a JS search index; treat them as one auth-gated, data-driven template '
    'rather than individual page builds, and migrate the data set rather than the rendered HTML.',
    'Card product/detail and card-lineup listing pages are the richest templates (carousel + tabs '
    '+ accordion + comparison tables) and are marketing-critical — build/QA these manually.',
    'Four distinct page "shells" exist (standard /ja/, AEM corporate_site, legacy /corporate '
    'microsite, legacy /en). Recommend consolidating header/footer/nav into one EDS implementation.',
    'Standalone microsites (identity-verification honninkakunin, entry-form application LPs, '
    'Club Online bumpers, cancellation flow) are outside the AEM shell and should be scoped '
    'individually or re-pointed to their owning systems.',
    'E-statement email templates (337) and SSI/partial fragments (145) are not website pages and '
    'should be removed from the page-count baseline used for migration estimation.',
    'Several legacy /corporate and /en URLs returned redirects or 404 during inspection — confirm '
    'the live, canonical URL set with the client before committing final counts.',
]:
    doc.add_paragraph(n, style='List Bullet')

doc.save(OUT)
print('Saved:', OUT, '| categories:', len(rows), '| distinct total:', total)
