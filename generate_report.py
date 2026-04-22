#!/usr/bin/env python3
"""Generate comprehensive site analysis Word document for LG Russia (lg.com/ru)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS_DIR = "/tmp/playwright/screenshots"

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, 'A50034')  # LG brand red

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if r % 2 == 0:
                set_cell_shading(cell, 'F5F5F5')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_screenshot(doc, filename, caption, width=5.5):
    """Add a screenshot image with caption."""
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ===================== COVER PAGE =====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('LG Russia Website Analysis')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(165, 0, 52)  # LG red

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Comprehensive Site Analysis & Migration Assessment')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Target Site: https://www.lg.com/ru\nDate: April 2026\nPrepared for: AEM Edge Delivery Services Migration')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===================== TABLE OF CONTENTS =====================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Templates Inventory',
        '3. Blocks / Components Catalog',
        '4. Page Counts by Template',
        '5. Integrations Analysis',
        '6. Complex Use Cases & Observations',
        '7. Migration Estimates',
        '8. Appendix: Screenshots'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===================== 1. EXECUTIVE SUMMARY =====================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive analysis of the LG Russia website (www.lg.com/ru) '
        'for migration to Adobe Experience Manager (AEM) Edge Delivery Services. The analysis covers '
        'all page templates, reusable blocks/components, third-party integrations, complex use cases, '
        'and migration effort estimates.'
    )
    doc.add_paragraph()

    doc.add_heading('Key Findings', level=2)
    findings = [
        '12 distinct page templates identified across consumer and B2B sections',
        '25 reusable blocks/components cataloged with complexity ratings',
        '800+ press release articles represent the bulk of content pages',
        '~500+ product detail pages across all categories',
        '~45 category/subcategory listing pages',
        'Heavy reliance on AJAX/dynamic content loading and lazy-loading patterns',
        'Multiple third-party integrations including Akamai CDN, Boomerang RUM, Google Analytics',
        'Russian social media integrations (VK, OK, Telegram, Yandex Zen)',
        'Virtual Showroom uses panoramic 360-degree technology requiring special handling',
        'B2B section operates as a semi-independent sub-site with different navigation and footer'
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ===================== 2. TEMPLATES INVENTORY =====================
    doc.add_heading('2. Templates Inventory', level=1)
    doc.add_paragraph(
        'The following table lists all unique page templates identified across the LG Russia website. '
        'Each template serves a distinct purpose and has its own layout structure, component composition, and complexity level.'
    )

    templates = [
        ['T01 - Homepage', 'Complex',
         'Full-width hero carousel (7+ slides), product category cards, promotional banners, magazine articles grid, newsletter signup, support contact block, promotional campaign sections. Heavy JS with lazy-loading.',
         'https://www.lg.com/ru'],
        ['T02 - Product Category (PLP)', 'Complex',
         'Filter sidebar (10+ filter groups with accordion), product card grid, sorting controls, pagination, "Most Viewed" carousel, FAQ accordion, comparison toolbar. AJAX-driven filtering.',
         'https://www.lg.com/ru/televisions\nhttps://www.lg.com/ru/washing-machines\nhttps://www.lg.com/ru/monitors'],
        ['T03 - Product Subcategory', 'Complex',
         'Similar to PLP but with technology-specific hero banner at top, promotional content blocks between product listings, feature highlight sections, technology comparison tables.',
         'https://www.lg.com/ru/oled-televisions\nhttps://www.lg.com/ru/qned-tvs\nhttps://www.lg.com/ru/ultragear-oled-monitors'],
        ['T04 - Product Detail Page (PDP)', 'Complex',
         'Product image gallery with zoom, model selector tabs, key specs summary, "Where to Buy" CTA, feature sections with full-width images, detailed specs accordion, comparison tool, related products carousel, support links.',
         'https://www.lg.com/ru/televisions/lg-65ua75009la'],
        ['T05 - Magazine Listing', 'Medium',
         'Article card grid (3 columns), category badges (Tips, Events, Solutions, FAQ), "Load More" pagination, featured article hero at top. Custom CMS template separate from main site.',
         'https://www.lg.com/ru/lg-magazine'],
        ['T06 - Magazine Article', 'Medium',
         'Article header with social sharing (VK, OK, Telegram), date, hero image, rich text body with table of contents sidebar, embedded product cards, related articles navigation (prev/next), tag links.',
         'https://www.lg.com/ru/lg-magazine/how-to/televizory-lg-oled-evo-ai-2025-goda-obzor-i-sravneniye-seriy-g5-c5-i-b5'],
        ['T07 - Press Center', 'Medium',
         'Tabbed interface (Press Releases / Press About Us), article list with date, title, and excerpt, pagination (5 pages visible), bottom cards for press contacts and media gallery.',
         'https://www.lg.com/ru/about-lg/press-and-media'],
        ['T08 - Support Hub', 'Complex',
         'Search bar with predictive suggestions, notification carousel, product category selector, 12-card support services grid, promotional banner cards (4), contact methods section (7 channels), floating action bar.',
         'https://www.lg.com/ru/support'],
        ['T09 - About LG / Corporate', 'Simple',
         'Hero banner, 4 content cards (Brand, Career, Info Center, User Survey), footer navigation. Minimal interactivity.',
         'https://www.lg.com/ru/about-lg\nhttps://www.lg.com/ru/about-lg/history\nhttps://www.lg.com/ru/about-lg/career'],
        ['T10 - B2B Landing Page', 'Complex',
         'Separate header/navigation from consumer site, hero carousel (5 slides), product mosaic tiles, category grid, contact section with B2B-specific phone/email, giveaway registration form with validation.',
         'https://www.lg.com/ru/business'],
        ['T11 - Landing / Campaign Pages', 'Medium',
         'Full-width lifestyle sections, product showcase blocks, embedded video thumbnails, progressive scroll reveals, CTA buttons. Used for brand campaigns (LG AI, LG Signature, LG ThinQ).',
         'https://www.lg.com/ru/lg-ai\nhttps://www.lg.com/ru/lg-signature\nhttps://www.lg.com/ru/lg-thinq'],
        ['T12 - Virtual Showroom', 'Complex',
         'Panoramic 360-degree room navigation, AR QR code, room selector (4 rooms), product simulator with filters, interactive product demos (drag-and-drop, temperature controls), product carousel.',
         'https://www.lg.com/ru/virtualshowroom'],
    ]

    add_table_with_style(
        doc,
        ['Template ID & Name', 'Complexity', 'Description', 'Reference URL(s)'],
        templates,
        col_widths=[4.5, 2, 7, 5]
    )

    doc.add_paragraph()
    add_screenshot(doc, '01-homepage-full.png', 'Figure 1: Homepage Template (T01)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '02-category-plp.png', 'Figure 2: Product Category Listing (T02)', 3.5)

    doc.add_page_break()
    add_screenshot(doc, '08-product-detail-page.png', 'Figure 3: Product Detail Page (T04)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '05-magazine-listing.png', 'Figure 4: Magazine Listing (T05)', 3.5)

    doc.add_page_break()
    add_screenshot(doc, '09-magazine-article.png', 'Figure 5: Magazine Article Detail (T06)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '06-press-releases.png', 'Figure 6: Press Center (T07)', 3.5)

    doc.add_page_break()
    add_screenshot(doc, '04-support-page.png', 'Figure 7: Support Hub (T08)', 3.5)
    doc.add_paragraph()
    add_screenshot(doc, '07-business-page.png', 'Figure 8: B2B Landing Page (T10)', 3.5)

    doc.add_page_break()

    # ===================== 3. BLOCKS / COMPONENTS CATALOG =====================
    doc.add_heading('3. Blocks / Components Catalog', level=1)
    doc.add_paragraph(
        'The following catalog identifies all reusable blocks and components across the LG Russia site. '
        'Where the content model is the same but the visual layout differs, design variations of the same block are noted rather than separate blocks.'
    )

    blocks = [
        ['B01 - Global Header', 'High',
         'Mega-menu navigation with multi-level dropdowns, brand sub-links (LG Signature, ThinQ), B2B/Consumer toggle, global search with predictive suggestions. Sticky on scroll. Different variant for B2B section.',
         'All pages'],
        ['B02 - Global Footer', 'Medium',
         'Multi-column sitemap links, social media icons (VK, OK, Telegram, YouTube, Yandex Zen), region/language selector, legal links, copyright. B2B variant has different link structure.',
         'All pages'],
        ['B03 - Hero Carousel', 'High',
         'Full-width image carousel with auto-play, pause/play controls, dot navigation, prev/next arrows. Supports responsive images (desktop/mobile variants). 5-7 slides typical. Multiple design variations: homepage (full-bleed), PLP (with text overlay), B2B (with CTA buttons).',
         'https://www.lg.com/ru\nhttps://www.lg.com/ru/business'],
        ['B04 - Product Card', 'High',
         '"New" badge, product image with size tabs, model name with copy button, key specs bullets (3-4), "Add to Cart" CTA, "Where to Buy" link, comparison checkbox, social share buttons (VK, OK, copy URL). Lazy-loaded.',
         'https://www.lg.com/ru/televisions\nhttps://www.lg.com/ru/monitors'],
        ['B05 - Filter Sidebar', 'High',
         'Accordion-style filter groups (10+ categories: type, screen size, resolution, gaming features, year, OS, refresh rate). Multi-select checkboxes, "Clear All" reset, live result count update. AJAX-driven.',
         'https://www.lg.com/ru/televisions\nhttps://www.lg.com/ru/washing-machines'],
        ['B06 - Product Image Gallery', 'High',
         'Main product image with zoom on hover, thumbnail strip navigation, size variant switcher tabs, 360-degree view toggle where available. Multiple image formats.',
         'All PDP pages'],
        ['B07 - Specs Accordion', 'Medium',
         'Expandable sections for detailed product specifications. Organized by category (General, Display, Audio, Connectivity, etc.). "Show All" toggle.',
         'All PDP pages'],
        ['B08 - Comparison Tool', 'High',
         'Floating bottom toolbar showing selected products count. Max product limit per category. Side-by-side spec comparison page. Add/remove products dynamically.',
         'All PLP/PDP pages'],
        ['B09 - Magazine Article Card', 'Low',
         'Thumbnail image (1280x960), category badge (color-coded: Tips=green, Events=blue, Solutions=purple), headline, excerpt, "Read More" link. Grid layout (3 columns).',
         'https://www.lg.com/ru/lg-magazine'],
        ['B10 - Notification Banner', 'Low',
         'Dismissible top banner with promo text and CTA button. Used for Virtual Showroom promo, cookie consent, browser compatibility warnings.',
         'Multiple pages'],
        ['B11 - Breadcrumb Navigation', 'Low',
         'Hierarchical path with slash separators. Links to parent pages. Bold current page.',
         'All interior pages'],
        ['B12 - Pagination', 'Low',
         'Numbered page buttons, Previous/Next controls, "Show More" variant for magazine. Two design variants: numbered (PLP, Press) and "Load More" (Magazine).',
         'PLP, Press, Magazine pages'],
        ['B13 - Social Share Bar', 'Low',
         'VKontakte, Odnoklassniki, Telegram share icons. Horizontal layout on article pages. Product card variant includes URL copy.',
         'Magazine articles, PDP pages'],
        ['B14 - Support Services Grid', 'Medium',
         '12 service cards in 2-column grid. Each has icon, title, description. Links to specific support functions (manuals, firmware, chat, repair, warranty, etc.).',
         'https://www.lg.com/ru/support'],
        ['B15 - Contact Methods Block', 'Medium',
         '7 contact channels with icons: Chat & Email, Phone, Survey, Telegram, Viber, WhatsApp, CEO Feedback. Horizontal layout.',
         'https://www.lg.com/ru/support'],
        ['B16 - Where to Buy Section', 'Medium',
         'Retailer links/buttons, store locator integration, online purchase options. Appears on PDP pages.',
         'All PDP pages'],
        ['B17 - FAQ Accordion', 'Low',
         'Expandable Q&A pairs. Used at bottom of category pages. Schema.org FAQ markup for SEO.',
         'Category listing pages'],
        ['B18 - Newsletter / Giveaway Form', 'Medium',
         'Email input, name field, phone field with validation, consent checkboxes (contest rules, privacy, data processing, marketing). Submit button with success/error messaging.',
         'Homepage, B2B landing'],
        ['B19 - Product Mosaic Grid', 'Medium',
         'Category tiles with product image, name, description, and hover interaction. "move" indicator for interactive tiles. Used on B2B landing and homepage.',
         'https://www.lg.com/ru/business'],
        ['B20 - Table of Contents Sidebar', 'Low',
         'Sticky sidebar with anchor links to article sections. Collapsible with "Expand" button. Appears on long-form magazine articles.',
         'Magazine article pages'],
        ['B21 - Related Products Carousel', 'Medium',
         'Horizontal scrolling product cards with left/right arrows. "Most Viewed" variant on PLP. "Related Models" variant on Magazine articles.',
         'PLP and Magazine article pages'],
        ['B22 - Promotional Banner Cards', 'Medium',
         '2x2 grid of large image cards with overlay text and CTA. Used for featured content like accessories, fraud warnings, branded service.',
         'https://www.lg.com/ru/support'],
        ['B23 - Video Thumbnail Grid', 'Medium',
         'YouTube video preview thumbnails in grid layout. Click to play in modal or new tab. Used on landing pages (LG AI, LG Signature).',
         'https://www.lg.com/ru/lg-ai\nhttps://www.lg.com/ru/lg-signature'],
        ['B24 - Floating Action Bar', 'Medium',
         'Fixed right-side panel with quick-access buttons: Contact, Email, Chat, Survey, Recently Viewed (with counter). Collapsible.',
         'Support and product pages'],
        ['B25 - Cookie Consent Banner', 'Low',
         'Bottom overlay with accept/decline options and cookie preferences link. e-privacy implementation.',
         'All pages (first visit)'],
    ]

    add_table_with_style(
        doc,
        ['Block ID & Name', 'Complexity', 'Description & Functionality', 'Reference URL(s)'],
        blocks,
        col_widths=[4, 2, 8, 4.5]
    )

    doc.add_page_break()

    # ===================== 4. PAGE COUNTS BY TEMPLATE =====================
    doc.add_heading('4. Page Counts by Template', level=1)
    doc.add_paragraph(
        'The following table provides estimated page counts for each template type, along with '
        'migration feasibility assessment. Counts are based on sitemap analysis, navigation structure, '
        'and crawl sampling.'
    )

    page_counts = [
        ['T01 - Homepage', '1', 'Manual', 'High customization, complex carousel logic, multiple promotional blocks with campaign-specific content. Requires careful block mapping.'],
        ['T02 - Product Category (PLP)', '~15', 'Semi-Automatic', 'Standardized template but heavy AJAX filtering needs reimplementation. Product data feeds must be mapped. Filter logic is complex.'],
        ['T03 - Product Subcategory', '~30', 'Semi-Automatic', 'Similar to PLP but with technology-specific marketing content that varies per subcategory. Hero banners need manual curation.'],
        ['T04 - Product Detail Page (PDP)', '~500+', 'Automatic (structure) / Manual (content)', 'Product data can be bulk-imported from PIM/feed. But interactive features (gallery, comparison, specs accordion) need EDS block development.'],
        ['T05 - Magazine Listing', '1', 'Manual', 'Custom CMS integration. Article data aggregation and category filtering need reimplementation.'],
        ['T06 - Magazine Article', '~60-80', 'Semi-Automatic', 'Rich text content can be migrated. But embedded product cards, table of contents, and related articles logic need manual configuration.'],
        ['T07 - Press Center', '1 (hub) + 800+ articles', 'Automatic (articles) / Manual (hub)', 'Press release articles are standardized text content ideal for bulk import. Hub page pagination/tabs need manual setup.'],
        ['T08 - Support Hub', '~15', 'Manual', 'Multiple interactive components (search, repair request, service center locator) require integration work. Sub-pages (manuals, warranty, etc.) are moderately complex.'],
        ['T09 - About / Corporate', '~5', 'Semi-Automatic', 'Mostly static content pages with simple layout. Minor interactive elements (survey link).'],
        ['T10 - B2B Landing Page', '~20', 'Manual', 'Separate navigation, different footer, B2B-specific product categories. Giveaway form with validation. Requires dedicated block variants.'],
        ['T11 - Landing / Campaign', '~5-8', 'Manual', 'Highly visual, custom layouts per campaign. Progressive scroll reveals, video integration. Each page is unique.'],
        ['T12 - Virtual Showroom', '1', 'Manual (Complex)', 'Panoramic 360-degree technology, AR integration, interactive product demos. Most complex page requires custom WebGL/iframe solution.'],
    ]

    add_table_with_style(
        doc,
        ['Template', 'Est. Page Count', 'Migration Type', 'Notes / Rationale'],
        page_counts,
        col_widths=[4, 2.5, 3, 9]
    )

    doc.add_paragraph()
    doc.add_heading('Summary Totals', level=2)

    summary_counts = [
        ['Automatic Migration (bulk import)', '~800+', 'Press releases, standardized article content'],
        ['Semi-Automatic (template + manual review)', '~600+', 'PDPs, magazine articles, subcategories, corporate pages'],
        ['Manual Migration (custom development)', '~50-70', 'Homepage, support hub, B2B, campaign pages, virtual showroom'],
        ['TOTAL ESTIMATED PAGES', '~1,450-1,500+', 'All content across consumer and B2B sections'],
    ]

    add_table_with_style(
        doc,
        ['Migration Category', 'Est. Pages', 'Includes'],
        summary_counts,
        col_widths=[5, 3, 10]
    )

    doc.add_page_break()

    # ===================== 5. INTEGRATIONS ANALYSIS =====================
    doc.add_heading('5. Integrations Analysis', level=1)
    doc.add_paragraph(
        'The following table documents all third-party integrations and embedded services identified on the LG Russia website.'
    )

    integrations = [
        ['Akamai CDN', 'CDN / Infrastructure', 'Medium',
         'Content delivery network for static assets (images, JS, CSS). Edge server endpoints detected. Needs CDN reconfiguration for EDS.',
         'All pages'],
        ['Boomerang Real User Monitoring (RUM)', 'API / Analytics', 'Medium',
         'Client-side performance monitoring with API key. Tracks page load metrics, network performance, user experience scores.',
         'All pages'],
        ['Google Tag Manager / Analytics', 'Embed / Analytics', 'Low',
         'dataLayer implementation with custom variables: pageType, siteType (B2C/B2B), category hierarchy, product data. Standard GTM container.',
         'All pages'],
        ['VKontakte (VK) Social API', 'API / Social', 'Low',
         'Share functionality via VK API. Social media links in footer. Used for product sharing and article distribution.',
         'Product cards, magazine articles, footer'],
        ['Odnoklassniki (OK) Social API', 'API / Social', 'Low',
         'Share functionality via OK API. Social sharing on articles and products.',
         'Product cards, magazine articles, footer'],
        ['Telegram Integration', 'API / Messaging', 'Low',
         'Share links, chatbot integration (LG_Electronics_RUS_bot), channel links. Multiple touchpoints across support and sharing.',
         'Footer, support page, article sharing'],
        ['WhatsApp Business API', 'API / Messaging', 'Low',
         'Direct messaging link for customer support channel. No embedded widget, simple URL redirect.',
         'https://www.lg.com/ru/support'],
        ['Viber Messaging', 'API / Messaging', 'Low',
         'Deep link protocol for customer support channel.',
         'https://www.lg.com/ru/support'],
        ['Yandex Zen', 'Embed / Social', 'Low',
         'Content syndication channel link. Footer social media link.',
         'Footer on all pages'],
        ['YouTube', 'Embed / Video', 'Low',
         'Channel link. Video thumbnails on landing pages. No inline video embeds detected on main content pages.',
         'Footer, LG AI landing page'],
        ['e-Privacy / Cookie Consent', 'Embed / Compliance', 'Medium',
         'Cookie preference management via e-privacy.min.js. Consent banner with accept/decline. Required for GDPR/Russian data protection compliance.',
         'All pages'],
        ['Schema.org Structured Data', 'Custom Code / SEO', 'Medium',
         'Corporation, WebSite, SearchAction, Product, FAQ schema markup. JSON-LD implementation for search engine optimization.',
         'Homepage, PDP, category pages'],
        ['LG Account System', 'API / Authentication', 'High',
         'User authentication for wishlist, recently viewed products, repair requests, inquiry tracking. getAccessToken API calls detected on every page.',
         'All pages (header account link)'],
        ['Product Comparison API', 'API / Custom', 'High',
         'Internal API for dynamic product comparison. Manages product selection across category pages. Session-based state.',
         'All PLP/PDP pages'],
        ['Service Center Locator', 'API / Custom', 'High',
         'Geolocation-based service center finder. Interactive map integration with address/phone details. Repair request submission system.',
         'https://www.lg.com/ru/support/locate-repair-center'],
        ['Virtual Showroom Engine', 'Embed / Custom', 'High',
         'Panoramic 360-degree navigation engine using static JPG panoramas. AR capability via QR codes. Interactive product demos.',
         'https://www.lg.com/ru/virtualshowroom'],
        ['LG Ethics Hotline', 'Embed / External', 'Low',
         'External link to ethics.lg.co.kr for Jeong-Do Management compliance reporting.',
         'Footer on all pages'],
    ]

    add_table_with_style(
        doc,
        ['Integration Name', 'Type', 'Complexity', 'Description', 'Reference URL(s)'],
        integrations,
        col_widths=[3.5, 2.5, 2, 7, 3.5]
    )

    doc.add_page_break()

    # ===================== 6. COMPLEX USE CASES =====================
    doc.add_heading('6. Complex Use Cases & Observations', level=1)
    doc.add_paragraph(
        'The following section identifies complex behaviors, edge cases, and functionality that require special attention during migration.'
    )

    complex_cases = [
        ['AJAX Product Filtering & Sorting', '~45 pages (all PLP/subcategory)', 'Product listing pages',
         'Real-time filter application via URL parameters, AJAX-driven product list updates without full page reload. 10+ filter categories with multi-select. IntersectionObserver-based lazy loading. Requires complete reimplementation as EDS blocks or client-side JS solution.',
         'Critical path - core shopping experience'],
        ['Product Comparison Tool', '~45+ pages (all PLP/PDP)', 'PLP and PDP pages',
         'Cross-page product selection persistence, floating comparison toolbar, side-by-side specification comparison view. Category-limited selections. Session-based state management.',
         'Needs custom EDS block with local storage or session management'],
        ['Virtual Showroom 360-degree Experience', '1 page (4 rooms)', 'https://www.lg.com/ru/virtualshowroom',
         'Panoramic navigation, AR QR codes for mobile, interactive product demos (washing machine fabric selection, AC temperature controls, refrigerator learning mode). Custom rendering engine.',
         'Most complex page. Consider iframe embed or progressive web app approach.'],
        ['B2B Sub-site Architecture', '~20+ pages', 'https://www.lg.com/ru/business/*',
         'Completely different header/navigation/footer from consumer site. Separate product catalog (commercial displays, HVAC, laundry equipment). "Inquiry to Buy" flow instead of "Where to Buy". Different support phone numbers and contact channels.',
         'Requires separate EDS template set or conditional rendering logic'],
        ['Magazine CMS Platform', '~60-80 articles + hub', 'https://www.lg.com/ru/lg-magazine/*',
         'Separate CMS from main site (different JS framework). Custom category system (Tips, Events, Solutions, FAQ). Table of contents generation, embedded product cards, prev/next navigation. Tags system.',
         'Content can be migrated but CMS features need EDS-native reimplementation'],
        ['User Account & Session Management', 'All pages (~1500)', 'Site-wide',
         'getAccessToken calls on every page load. Recently viewed products tracking, wishlist functionality, repair request status tracking, inquiry status checking. Session-dependent features.',
         'OAuth/token integration required. Consider which features to maintain vs. simplify.'],
        ['Support Portal Interactive Features', '~15 pages', 'https://www.lg.com/ru/support/*',
         'Predictive search for model numbers, repair request submission forms, service center geolocation, warranty lookup, firmware download by model. Multiple API endpoints.',
         'High integration complexity. Each support sub-feature needs separate API mapping.'],
        ['Responsive Image Strategy', 'All pages', 'Site-wide',
         'Different image sources for desktop vs. mobile breakpoints. Akamai image optimization. Product images with multiple size variants. Lazy-loading via IntersectionObserver.',
         'EDS image optimization pipeline needed. Ensure CDN compatibility.'],
        ['Multi-channel Messaging Support', '~15 pages', 'Support section',
         'WhatsApp, Telegram (bot), Viber deep links. Online chat widget. Multiple contact channels with conditional display based on business hours (8:00-22:00, 365 days).',
         'Deep links are simple but chat/bot integrations may need platform-specific setup'],
        ['Giveaway / Registration Forms', '~2 pages', 'Homepage, B2B landing',
         'Form with name, phone, email validation. Contest rules acceptance. Privacy/data consent checkboxes. Success/error/duplicate subscriber messaging. Currently "inactive" state handling.',
         'Form block with validation and API submission needed. Russian data protection compliance.'],
    ]

    add_table_with_style(
        doc,
        ['Use Case', 'Instances / Scope', 'Where Found', 'Description', 'Why Complex / Migration Impact'],
        complex_cases,
        col_widths=[3.5, 2.5, 3, 5, 4.5]
    )

    doc.add_page_break()

    # ===================== 7. MIGRATION ESTIMATES =====================
    doc.add_heading('7. Migration Estimates', level=1)
    doc.add_paragraph(
        'The following estimates cover the full migration of lg.com/ru to AEM Edge Delivery Services. '
        'Estimates assume a team of 2-3 EDS developers, 1 content migration specialist, and 1 QA engineer.'
    )

    doc.add_heading('7.1 Effort Breakdown by Phase', level=2)

    effort = [
        ['Phase 1: Discovery & Design System', '', '', ''],
        ['  Design token extraction (colors, fonts, spacing)', '2-3 days', 'Semi-Automatic', 'Extract CSS custom properties from existing site'],
        ['  Template architecture design', '3-4 days', 'Manual', 'Map 12 templates to EDS page structures'],
        ['  Block specification & design', '4-5 days', 'Manual', 'Define 25 block specs with content models'],
        ['', '', '', ''],
        ['Phase 2: Block Development', '', '', ''],
        ['  Global Header (incl. mega-menu)', '4-5 days', 'Manual', 'Complex mega-menu with B2B variant'],
        ['  Global Footer (2 variants)', '1-2 days', 'Manual', 'Consumer + B2B variants'],
        ['  Hero Carousel (3 variants)', '3-4 days', 'Manual', 'Auto-play, responsive images, multiple layouts'],
        ['  Product Card block', '3-4 days', 'Manual', 'Multiple states, lazy loading, comparison checkbox'],
        ['  Filter Sidebar block', '5-7 days', 'Manual', 'Most complex block with AJAX filtering, multi-select'],
        ['  Product Image Gallery', '3-4 days', 'Manual', 'Zoom, thumbnails, variant switcher'],
        ['  Specs Accordion', '1-2 days', 'Manual', 'Expandable sections with schema'],
        ['  Comparison Tool', '4-5 days', 'Manual', 'Cross-page state, floating toolbar'],
        ['  Magazine Article Card', '1 day', 'Manual', 'Simple card with category badge'],
        ['  Support Services Grid', '1-2 days', 'Manual', 'Icon cards with links'],
        ['  Contact Methods Block', '1 day', 'Manual', 'Multi-channel contact display'],
        ['  FAQ Accordion', '1 day', 'Manual', 'Schema.org FAQ markup'],
        ['  Newsletter / Form blocks', '2-3 days', 'Manual', 'Validation, consent, API submission'],
        ['  Remaining blocks (10+)', '8-10 days', 'Manual', 'Pagination, breadcrumbs, ToC sidebar, etc.'],
        ['', '', '', ''],
        ['Phase 3: Content Migration', '', '', ''],
        ['  Press releases (~800 articles)', '3-5 days', 'Automatic', 'Bulk import of standardized content'],
        ['  Magazine articles (~60-80)', '3-4 days', 'Semi-Automatic', 'Content import + manual review of embedded products'],
        ['  Product pages (~500+)', '5-7 days', 'Semi-Automatic', 'PIM data mapping + template application'],
        ['  Category/subcategory pages (~45)', '3-4 days', 'Semi-Automatic', 'Template + manual hero/promo content'],
        ['  Support pages (~15)', '3-4 days', 'Manual', 'API integration mapping required'],
        ['  About/Corporate (~5)', '1-2 days', 'Manual', 'Simple content migration'],
        ['  B2B section (~20)', '3-4 days', 'Manual', 'Separate template set, product data'],
        ['  Campaign/Landing pages (~5-8)', '3-5 days', 'Manual', 'Custom layouts per page'],
        ['  Virtual Showroom', '5-7 days', 'Manual', 'Custom 360-degree solution development'],
        ['  Homepage', '2-3 days', 'Manual', 'Complex assembly of multiple blocks'],
        ['', '', '', ''],
        ['Phase 4: Integration Setup', '', '', ''],
        ['  Analytics (GTM/Boomerang)', '2-3 days', 'Manual', 'Tag configuration and dataLayer mapping'],
        ['  Social media integrations', '1-2 days', 'Manual', 'VK, OK, Telegram share APIs'],
        ['  User authentication', '3-5 days', 'Manual', 'Account system, session management'],
        ['  Support APIs (search, repair, locator)', '5-7 days', 'Manual', 'Multiple API endpoints'],
        ['  Cookie consent / e-privacy', '1-2 days', 'Manual', 'Consent management setup'],
        ['', '', '', ''],
        ['Phase 5: QA & Testing', '', '', ''],
        ['  Template/block visual QA', '5-7 days', 'Manual', 'Cross-browser, responsive testing'],
        ['  Content validation', '3-5 days', 'Semi-Automatic', 'Spot-check migrated content accuracy'],
        ['  Integration testing', '3-4 days', 'Manual', 'API endpoints, analytics verification'],
        ['  Performance testing', '2-3 days', 'Manual', 'Lighthouse, Core Web Vitals'],
        ['  UAT support', '3-5 days', 'Manual', 'Stakeholder review and feedback cycles'],
    ]

    add_table_with_style(
        doc,
        ['Task / Activity', 'Estimated Duration', 'Migration Type', 'Notes'],
        effort,
        col_widths=[6, 3, 3, 6.5]
    )

    doc.add_paragraph()
    doc.add_heading('7.2 Total Estimated Schedule', level=2)

    schedule = [
        ['Phase 1: Discovery & Design System', '2 weeks', '80-96 hrs'],
        ['Phase 2: Block Development', '6-8 weeks', '320-480 hrs'],
        ['Phase 3: Content Migration', '4-5 weeks', '200-280 hrs'],
        ['Phase 4: Integration Setup', '2-3 weeks', '96-152 hrs'],
        ['Phase 5: QA & Testing', '3-4 weeks', '128-192 hrs'],
        ['Buffer / Contingency (15%)', '2-3 weeks', '120-180 hrs'],
        ['TOTAL', '19-26 weeks (~5-6 months)', '944-1,380 hrs'],
    ]

    add_table_with_style(
        doc,
        ['Phase', 'Duration', 'Estimated Hours'],
        schedule,
        col_widths=[7, 5, 5]
    )

    doc.add_paragraph()
    doc.add_heading('7.3 Cost Estimate Assumptions', level=2)
    doc.add_paragraph(
        'The following cost assumptions are based on typical AEM Edge Delivery Services implementation rates:'
    )
    cost_items = [
        'Team size: 2-3 EDS developers, 1 content specialist, 1 QA engineer',
        'Parallel execution: Phases 2-4 can overlap with phased content migration',
        'With parallel execution, wall-clock time reduces to approximately 14-18 weeks',
        'Key risks: Virtual Showroom complexity, product data feed availability, API documentation access',
        'Recommendation: Phase the migration starting with Magazine/Press (quick wins), then Categories/PDPs, then complex pages last',
    ]
    for item in cost_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('7.4 Recommended Migration Approach', level=2)

    approach = [
        ['Wave 1 (Weeks 1-6)', 'Foundation + Quick Wins',
         'Design system extraction, global header/footer blocks, press release bulk import (~800 pages), magazine content migration (~80 pages). Delivers ~880 pages.'],
        ['Wave 2 (Weeks 5-12)', 'Product Experience',
         'Product card, filter sidebar, PDP blocks, product data import (~500 PDPs), category/subcategory pages (~45). Core shopping experience. Delivers ~545 pages.'],
        ['Wave 3 (Weeks 10-16)', 'Support & Integrations',
         'Support hub blocks, API integrations (search, repair, locator), user account features, B2B section (~20 pages). Delivers ~35 pages.'],
        ['Wave 4 (Weeks 14-20)', 'Premium & Custom',
         'Campaign landing pages, Virtual Showroom, LG Signature/AI/ThinQ brand pages, homepage. Delivers ~15 pages.'],
        ['Wave 5 (Weeks 18-22)', 'QA, UAT & Launch',
         'Comprehensive testing, stakeholder review, performance optimization, analytics verification, go-live preparation.'],
    ]

    add_table_with_style(
        doc,
        ['Wave / Timeline', 'Focus Area', 'Scope & Deliverables'],
        approach,
        col_widths=[3.5, 3.5, 11.5]
    )

    doc.add_page_break()

    # ===================== 8. APPENDIX =====================
    doc.add_heading('8. Appendix: Additional Screenshots', level=1)

    add_screenshot(doc, '01-homepage.png', 'Figure A1: Homepage - Above the Fold', 5)
    doc.add_paragraph()
    add_screenshot(doc, '03-subcategory-plp.png', 'Figure A2: Subcategory Page (OLED TVs) with Tech Marketing Content', 3.5)

    # Save
    output_path = '/workspace/LG_Russia_Site_Analysis_Report.docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')

if __name__ == '__main__':
    main()
