#!/usr/bin/env python3
"""Generate EDS migration analysis reports for the two Sumitomo Mitsui Trust Club
credit-card sites:
  - www.diners.co.jp   (Diners Club Japan)
  - www.sumitclub.jp   (TRUST CLUB)

Produces one Word document per site. Page counts are derived from the live
sitemaps (diners: 1,765 pages; sumitclub: 439 pages) classified into templates.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

TABLE_STYLE = 'Medium Shading 1 Accent 1'


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------
def new_doc():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    return doc


def add_table(doc, data, widths=None):
    cols = len(data[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(data[0]):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in data[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def title_page(doc, site_title, site_url, totals):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_heading(site_title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading('Adobe Edge Delivery Services (EDS) Migration Analysis', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Date: June 2026\n')
    run.bold = True
    info.add_run(f'Source URL: {site_url}\n')
    info.add_run(f'Total pages analyzed (from sitemap): {totals}\n')
    info.add_run('Operator: Sumitomo Mitsui Trust Club Co., Ltd.\n')
    info.add_run('Source platform: Adobe Experience Manager (AEM Sites)')
    doc.add_page_break()


def toc(doc):
    doc.add_heading('Table of Contents', level=1)
    for item in [
        '1. Executive Summary',
        '2. Templates Inventory',
        '3. Blocks / Components Catalog',
        '4. Page Counts by Template & Migration Approach',
        '5. Integrations Analysis',
        '6. Migration Estimates',
        '7. Assumptions & Notes',
    ]:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()


def build_report(cfg, out_path):
    doc = new_doc()
    title_page(doc, cfg['title'], cfg['url'], cfg['total_pages'])
    toc(doc)

    # 1. Executive summary -------------------------------------------------
    doc.add_heading('1. Executive Summary', level=1)
    for para in cfg['exec_summary']:
        doc.add_paragraph(para)
    add_table(doc, [
        ['Metric', 'Value'],
        ['Total pages (sitemap)', str(cfg['total_pages'])],
        ['Unique page templates', str(len(cfg['templates']) - 1)],
        ['Reusable blocks / components', str(len(cfg['blocks']) - 1)],
        ['Auto-migratable pages', cfg['auto_pages']],
        ['Manual / assisted pages', cfg['manual_pages']],
        ['Estimated total effort', cfg['effort_total']],
        ['Estimated schedule', cfg['schedule']],
    ])
    doc.add_page_break()

    # 2. Templates inventory ----------------------------------------------
    doc.add_heading('2. Templates Inventory', level=1)
    doc.add_paragraph(cfg['templates_intro'])
    add_table(doc, cfg['templates'])
    doc.add_paragraph()
    doc.add_paragraph(cfg['templates_note'])
    doc.add_page_break()

    # 3. Blocks catalog ----------------------------------------------------
    doc.add_heading('3. Blocks / Components Catalog', level=1)
    doc.add_paragraph(cfg['blocks_intro'])
    add_table(doc, cfg['blocks'])
    doc.add_paragraph()
    doc.add_heading('3.1 Reconciliation with Team-Identified Block List', level=2)
    doc.add_paragraph(cfg['reconcile_intro'])
    add_table(doc, cfg['reconcile'])
    doc.add_page_break()

    # 4. Page counts -------------------------------------------------------
    doc.add_heading('4. Page Counts by Template & Migration Approach', level=1)
    doc.add_paragraph(cfg['counts_intro'])
    add_table(doc, cfg['counts'])
    doc.add_paragraph()
    doc.add_paragraph(
        'Migration approach key — "Automated": pages with standardized structure that the '
        'import pipeline (block parsers + page transformers) can convert in bulk with spot-check QA. '
        '"Assisted": bulk-imported but require per-page review/cleanup. '
        '"Manual": custom layout, dynamic/system-driven, or external-integration pages that must be '
        'rebuilt or re-pointed by hand.'
    )
    doc.add_page_break()

    # 5. Integrations ------------------------------------------------------
    doc.add_heading('5. Integrations Analysis', level=1)
    doc.add_paragraph(cfg['integrations_intro'])
    add_table(doc, cfg['integrations'])
    doc.add_page_break()

    # 6. Estimates ---------------------------------------------------------
    doc.add_heading('6. Migration Estimates', level=1)
    doc.add_paragraph(cfg['estimates_intro'])

    doc.add_heading('6.1 Effort by Work-stream', level=2)
    add_table(doc, cfg['effort'])
    doc.add_paragraph()

    doc.add_heading('6.2 Content Migration Effort by Template', level=2)
    add_table(doc, cfg['content_effort'])
    doc.add_paragraph()

    doc.add_heading('6.3 Timeline / Schedule', level=2)
    add_table(doc, cfg['timeline'])
    doc.add_paragraph()

    doc.add_heading('6.4 Cost Summary', level=2)
    add_table(doc, cfg['cost'])
    doc.add_paragraph()
    doc.add_paragraph(cfg['cost_note'])
    doc.add_page_break()

    # 7. Assumptions -------------------------------------------------------
    doc.add_heading('7. Assumptions & Notes', level=1)
    for a in cfg['assumptions']:
        doc.add_paragraph(a, style='List Bullet')

    doc.save(out_path)
    print(f'Saved: {out_path}')


# ===========================================================================
# Shared block catalog (both sites run on the same SMTC AEM platform)
# ===========================================================================
COMMON_BLOCKS_HEADER = ['Block / Component', 'Variants', 'Complexity',
                        'Description & Behaviour']


# ===========================================================================
# DINERS configuration
# ===========================================================================
diners = {
    'title': 'Diners Club Japan (www.diners.co.jp)',
    'url': 'https://www.diners.co.jp/',
    'total_pages': 1765,
    'auto_pages': '≈ 1,610 (91%)',
    'manual_pages': '≈ 155 (9%)',
    'effort_total': '≈ 92–110 person-days',
    'schedule': '≈ 13–15 weeks (3–3.5 months)',
    'exec_summary': [
        'www.diners.co.jp is the consumer website for Diners Club Japan, operated by '
        'Sumitomo Mitsui Trust Club Co., Ltd. (SMTC). It is built on Adobe Experience '
        'Manager (AEM Sites) using a mix of WCM Core Components (cmp-*) and bespoke custom '
        'components (CCM### family), with the global header and footer delivered as '
        'Experience Fragments shared with its sister site www.sumitclub.jp.',
        'The live sitemap exposes 1,765 published pages. The content is dominated by three '
        'high-volume, highly-templated families: the editorial "Magazine" (601 article '
        'pages), "Benefit / Merchant" detail pages (470), and dated "Press / News" articles '
        '(258). These standardized families are excellent candidates for automated bulk '
        'import into EDS, which is why ~91% of the estate is auto-migratable.',
        'The migration is low-to-medium risk overall. The main complexity is not page volume '
        'but the surrounding integrations — site search, the Club Online member system, the '
        'Helpfeel FAQ, the external card-application forms, and heavy ad/analytics tagging — '
        'all of which are external systems that should be re-pointed rather than rebuilt.',
    ],
    'templates_intro': (
        'Nine unique page templates were identified across the 1,765 pages. Templates are '
        'grouped by structural layout and content pattern rather than by URL folder, so a '
        'single template typically spans many sections of the site.'
    ),
    'templates': [
        ['Template', 'Complexity', 'Reasoning', 'Reference URLs'],
        ['Homepage', 'High',
         'Hero carousel (CCM020), audience-segment navigation grid, "Important notices" '
         'ribbon, dual news feeds (お知らせ / ご確認ください), promo tiles. One-off layout.',
         'https://www.diners.co.jp/ja/index.html'],
        ['Article (Magazine)', 'Low',
         'Standardized editorial article: breadcrumb, title, hero image, rich-text body, '
         'inline images, related-article list, category tags. Highly repetitive (601 pages).',
         'https://www.diners.co.jp/ja/magazine/article/shopping/watch_jewelry/tiffany.html'],
        ['Benefit / Merchant detail', 'Medium',
         'Bespoke template: image gallery, benefit-summary tables, store-info table, map/'
         'access text, CTA. Consistent structure but data-dense (470 pages).',
         'https://www.diners.co.jp/ja/benefit/detail.R1882.html'],
        ['Article / News detail', 'Low',
         'Dated press-release / event-report article: date, category, title, rich text, '
         'attachments. Very standardized (258 pages).',
         'https://www.diners.co.jp/ja/press/inf_20260601.html\n'
         'https://www.diners.co.jp/ja/event/report/e00160.html'],
        ['Category / Service / How-to', 'Medium',
         'Core-Component pages (travel, gourmet, golf, usage, point, payment, insurance, '
         'corporate, merchant): page heading, side sub-nav (CCM027), image+text, step '
         'blocks, accordion, CTA, tables. Layout varies per section (374 pages).',
         'https://www.diners.co.jp/ja/travel/hotel.html\n'
         'https://www.diners.co.jp/ja/usage/applepay.html'],
        ['Product / Card detail', 'Medium',
         'Card-spec layout: hero card image, feature grid, fee/benefit tables, comparison, '
         'apply CTA to external entry form (43 pages).',
         'https://www.diners.co.jp/ja/cardlineup/dinersclubcard.html'],
        ['Listing / Index', 'Medium',
         'Section landing/listing pages with filters, card grids and load-more news lists '
         '(magazine library, press index, benefit index, card lineup).',
         'https://www.diners.co.jp/ja/magazine/library.html\n'
         'https://www.diners.co.jp/ja/press.html'],
        ['Legal / Policy text', 'Low',
         'Long-form text: headings, paragraphs, tables, link lists. Privacy, terms, small '
         'print, customer-response policy (13 pages).',
         'https://www.diners.co.jp/ja/privacy.html\n'
         'https://www.diners.co.jp/ja/tc.html'],
        ['Contact / Support', 'Medium',
         'Accordion of inquiry types + gateway buttons that launch external form widget / '
         'Helpfeel FAQ / phone directory (1 hub page).',
         'https://www.diners.co.jp/ja/contact.html'],
        ['Sitemap', 'Low',
         'Single hierarchical link-list page.',
         'https://www.diners.co.jp/ja/sitemap.html'],
    ],
    'templates_note': (
        'Note: header and footer are AEM Experience Fragments shared with www.sumitclub.jp '
        'and contain cross-domain links and Adobe Experience Cloud (adobe_mc) ID-sync '
        'parameters. These should become a single EDS nav/footer block, and the cross-domain '
        'parameters can be dropped on EDS.'
    ),
    'blocks_intro': (
        'The following 24 reusable blocks were identified across the site. Per the brief, '
        'visually different layouts that share the same content model are captured as '
        '"variants" of one block rather than as new blocks (e.g. card grid vs. news list).'
    ),
    'blocks': [
        COMMON_BLOCKS_HEADER,
        ['Global Header / Nav', '1', 'High',
         'Logo, utility links (Club Online, lost-card, search), site search box, corporate-'
         'site link, primary nav. Experience Fragment shared with sumitclub.jp.'],
        ['Global Footer', '1', 'Medium',
         'Brand/partner logos, Remote Operator link, corporate link, copyright. Experience '
         'Fragment.'],
        ['Breadcrumb', '1', 'Low', 'Hierarchical path navigation (CCM013).'],
        ['Side / Secondary Nav', '1', 'Low', 'In-section sub-navigation menu (CCM027).'],
        ['Hero / Page Heading', '2', 'Low',
         'Page title with optional background/hero image. Variant: card-image hero on product pages.'],
        ['Carousel / Banner', '1', 'Medium',
         'Auto-advancing banner slider with prev/next + dots (CCM020 / slick).'],
        ['Card Grid / Listing', '3', 'Medium',
         'Grid of linked cards (image+title+text). Variants: benefit grid, magazine grid, '
         'product/card lineup grid.'],
        ['News / Press List', '2', 'Medium',
         'Dated list with category badge + title + "load more". Variants: homepage feed, '
         'full press/notice index.'],
        ['Related Items List', '1', 'Low', 'Related-article / related-card list at page foot.'],
        ['Rich Text', '1', 'Low', 'Body copy with headings, lists, links (CCM006).'],
        ['Raw HTML Embed', '1', 'Medium', 'Arbitrary embedded HTML/markup block (CCM005).'],
        ['Data Table', '2', 'Low',
         'Fee/benefit/spec tables. Variants: spec table, comparison table.'],
        ['Accordion / Toggle', '1', 'Low', 'Expand/collapse FAQ or content sections.'],
        ['Tabs', '1', 'Medium', 'Tabbed content switcher.'],
        ['Image + Text', '2', 'Low',
         'Media-and-copy block. Variants: image-left, image-right.'],
        ['Step / Process Block', '1', 'Low', 'Numbered how-to / step sequence.'],
        ['CTA Button', '2', 'Low',
         'Primary call-to-action (CCM009). Variants: internal link, external/apply link.'],
        ['Link List', '2', 'Low',
         'Grouped lists of links (CCM008). Variants: inline list, sitemap tree.'],
        ['Image Gallery / Thumbnails', '1', 'Medium',
         'Thumbnail gallery on benefit/merchant pages.'],
        ['Social Share', '1', 'Low', 'Share buttons (JS window.open) on articles.'],
        ['Notice / Ribbon', '1', 'Low', 'Highlighted important-notice strip.'],
        ['Category Tag / Badge', '1', 'Low', 'Article category labels.'],
        ['Video Embed', '1', 'Medium', 'YouTube embeds (brand movie, articles).'],
        ['Search Box', '1', 'Medium',
         'Header site-search form (GET to search.diners.co.jp). External search engine.'],
    ],
    'reconcile_intro': (
        'The team\'s Excel block/integration sheet (Sheet1 = blocks, Sheet2 = integrations) '
        'was not available to this analysis at the time of writing. The table below lists the '
        'blocks identified from the live site so they can be diffed against the team list. '
        'Please confirm the "vs. Team List" column; items marked "Confirm" are additions or '
        'consolidations this analysis recommends.'
    ),
    'reconcile': [
        ['Recommendation', 'Block(s)', 'Rationale'],
        ['Consolidate as variants',
         'Card Grid (benefit / magazine / product); News List (homepage / index); '
         'Image+Text (left / right); Hero (standard / card)',
         'Same content model, different visual layout — build one block with variants '
         'instead of separate blocks, reducing block count and dev effort.'],
        ['Confirm / likely additions',
         'Raw HTML Embed (CCM005), Side/Secondary Nav (CCM027), Step/Process block, '
         'Image Gallery, Notice ribbon, Social Share',
         'Present on the live site; verify these exist in the team list, as they are easy '
         'to overlook but affect parser scope.'],
        ['Treat as integration, not block',
         'Search Box, Login/Club Online button, FAQ launcher',
         'These are gateways to external systems; scope them in the Integrations section, '
         'not as content blocks.'],
    ],
    'counts_intro': (
        'Page counts below are derived by classifying all 1,765 sitemap URLs into the nine '
        'templates from Section 2.'
    ),
    'counts': [
        ['Template', 'Pages', '% of site', 'Migration approach', 'Rationale'],
        ['Article (Magazine)', '601', '34.1%', 'Automated',
         'Uniform article structure; bulk import with spot-check QA.'],
        ['Benefit / Merchant detail', '470', '26.6%', 'Assisted',
         'Consistent template but data-dense tables/galleries; bulk import + per-page review.'],
        ['Category / Service / How-to', '374', '21.2%', 'Assisted',
         'Core-Component pages with per-section layout variation; bulk import + cleanup.'],
        ['Article / News detail', '258', '14.6%', 'Automated',
         'Highly standardized dated press/event articles.'],
        ['Product / Card detail', '43', '2.4%', 'Manual',
         'Marketing-critical card pages; rebuild for pixel-fidelity + external apply CTA.'],
        ['Legal / Policy text', '13', '0.7%', 'Automated',
         'Plain long-form text + tables.'],
        ['Listing / Index', '4', '0.2%', 'Manual',
         'Filtering / load-more behaviour needs an EDS index + query.'],
        ['Contact / Support', '1', '0.1%', 'Manual',
         'External form widget + FAQ launcher wiring.'],
        ['Sitemap', '1', '0.1%', 'Automated', 'Single link-list page.'],
        ['TOTAL', '1,765', '100%', '—', '—'],
    ],
    'integrations_intro': (
        'The site embeds several third-party / external systems. Most should be re-pointed '
        'or re-embedded in EDS rather than reimplemented.'
    ),
    'integrations': [
        ['Integration', 'Purpose', 'Migration handling'],
        ['Site Search (search.diners.co.jp)', 'On-site search', 'Re-point header search form, or replace with EDS/Edge search.'],
        ['Club Online (sumitclub.jp/JPCRD)', 'Member login / self-service', 'External system — link out; out of EDS scope.'],
        ['Helpfeel', 'FAQ / help knowledge base', 'External — keep as outbound link / embed.'],
        ['Card application forms (entry_form/lp)', 'New-card applications', 'External landing pages — link out.'],
        ['External inquiry form widget', 'General contact form', 'JS-injected widget — re-embed or replace with EDS form.'],
        ['YouTube', 'Brand / article video', 'Re-embed via EDS video block.'],
        ['e-book / library viewer', 'Magazine digital viewer', 'External viewer — link/embed.'],
        ['Adobe Analytics / Target + adobe_mc', 'Analytics & personalization, cross-domain ID', 'Re-tag via EDS; drop cross-domain params.'],
        ['Meta Pixel / Logly / ad trackers', 'Advertising & remarketing', 'Re-add required tags only; many can be retired.'],
        ['Remote Operator (sumitclub.rmop.jp)', 'Co-browse support', 'External — link out.'],
    ],
    'estimates_intro': (
        'Estimates assume an automated EDS import pipeline (block parsers + page '
        'transformers) is built once and reused across all standardized templates, with a '
        'blended rate and a small migration team (2–3 engineers + 1 QA). Page volume is high '
        'but highly repetitive, so per-page cost is low for the article/news families.'
    ),
    'effort': [
        ['Work-stream', 'Effort (person-days)', 'Notes'],
        ['Discovery & template design', '6–8', 'Confirm templates, design EDS block models.'],
        ['Block development (24 blocks, with variants)', '22–26',
         'Build + style blocks; shared header/footer; variant consolidation reduces count.'],
        ['Import pipeline (parsers + transformers)', '12–15',
         'Per-template parsers for magazine, benefit, press/event, category, product.'],
        ['Automated content migration (≈1,610 pages)', '10–14',
         'Bulk runs + reruns; mostly compute, light human time.'],
        ['Manual / assisted pages (≈155 pages)', '14–18',
         'Card pages, listings, contact, benefit/category cleanup.'],
        ['Integrations re-wiring', '8–10', 'Search, forms, video, analytics, login links.'],
        ['QA & Testing', '14–16', 'Visual QA, link checks, responsive, cross-browser.'],
        ['PM / coordination / UAT support', '6–8', 'Across the engagement.'],
        ['TOTAL', '92–110', '≈ 13–15 weeks elapsed with a 2–3 person team.'],
    ],
    'content_effort': [
        ['Template', 'Pages', 'Approach', 'Effort'],
        ['Article (Magazine)', '601', 'Automated', '3–4 days (bulk + QA sampling)'],
        ['Benefit / Merchant detail', '470', 'Assisted', '6–8 days (import + review)'],
        ['Category / Service / How-to', '374', 'Assisted', '5–6 days (import + cleanup)'],
        ['Article / News detail', '258', 'Automated', '1.5–2 days'],
        ['Product / Card detail', '43', 'Manual', '5–6 days (high fidelity)'],
        ['Legal / Policy text', '13', 'Automated', '0.5 day'],
        ['Listing / Index', '4', 'Manual', '2–3 days (index + query)'],
        ['Contact / Support', '1', 'Manual', '1–2 days (form wiring)'],
        ['Sitemap', '1', 'Automated', '0.25 day'],
    ],
    'timeline': [
        ['Phase', 'Duration', 'Key outputs'],
        ['Phase 1 — Discovery & design', 'Weeks 1–2', 'Template catalog, block models, pipeline design.'],
        ['Phase 2 — Block & pipeline build', 'Weeks 3–6', 'All blocks + import parsers/transformers ready.'],
        ['Phase 3 — Bulk migration', 'Weeks 6–9', 'Automated import of ~1,610 standardized pages.'],
        ['Phase 4 — Manual pages & integrations', 'Weeks 9–12', 'Card pages, listings, contact, integrations.'],
        ['Phase 5 — QA, UAT & launch', 'Weeks 12–15', 'Visual QA, fixes, go-live.'],
    ],
    'cost': [
        ['Scenario', 'Effort (person-days)', 'Indicative range*'],
        ['Low (optimistic)', '92', '92 × blended day-rate'],
        ['High (conservative)', '110', '110 × blended day-rate'],
    ],
    'cost_note': (
        '*Cost is expressed in person-days; apply your blended daily rate to convert to '
        'currency. As an illustration, at a representative offshore/near-shore blended rate '
        'of USD 500–700/day, the engagement falls roughly in the USD 46k–77k range. Replace '
        'with your contracted rate for a firm number.'
    ),
    'assumptions': [
        'Page counts are taken from the live sitemap (https://www.diners.co.jp/ja/Sitemap.xml) '
        'on the analysis date; 1,765 published pages, single Japanese locale.',
        'Header and footer are shared Experience Fragments with www.sumitclub.jp; building '
        'them once can be amortized if both sites are migrated together.',
        'External systems (Club Online member portal, Helpfeel FAQ, card-application forms, '
        'site search) remain on their current platforms and are linked/embedded, not rebuilt.',
        'Ad-tech and analytics tags are rationalized during migration; only business-required '
        'tags are re-added.',
        'The team\'s Excel block/integration list was not available at writing time; Section '
        '3.1 should be reconciled against it before finalizing scope.',
        'Estimates assume a 2–3 engineer + 1 QA team and reuse of the EDS import accelerator '
        'tooling; they exclude content rewriting, new design, translation, and post-launch '
        'support.',
    ],
}


# ===========================================================================
# SUMITCLUB configuration
# ===========================================================================
sumitclub = {
    'title': 'TRUST CLUB / Sumitomo Mitsui Trust Club (www.sumitclub.jp)',
    'url': 'https://www.sumitclub.jp/',
    'total_pages': 439,
    'auto_pages': '≈ 300 (68%)',
    'manual_pages': '≈ 139 (32%)',
    'effort_total': '≈ 60–74 person-days',
    'schedule': '≈ 9–11 weeks (2–2.5 months)',
    'exec_summary': [
        'www.sumitclub.jp is the TRUST CLUB website operated by Sumitomo Mitsui Trust Club '
        'Co., Ltd. (SMTC), the same company behind www.diners.co.jp. It runs on the same '
        'Adobe Experience Manager platform and shares header/footer Experience Fragments and '
        'the Club Online member system with its sister site.',
        'The live sitemap exposes 439 published pages across both the Japanese (/ja/) and '
        'English (/en/) locales plus a no-locale corporate area. Unlike diners.co.jp, the '
        'estate is smaller but structurally more heterogeneous: it contains four distinct '
        'page "shells" (standard /ja/, corporate-site, minimal application/help microsites, '
        'and a legacy /en/ shell) plus 40 login/member-system pages and a legacy table-based '
        'insurance microsite.',
        'Because of this shell fragmentation, a larger share of pages (~32%) needs manual or '
        'assisted handling, and several page families (login pages, verification microsites, '
        'system error pages) are functional rather than editorial and may be out of scope or '
        're-pointed to external systems. The standardized usage/how-to, notice, and corporate '
        'content remains a good fit for automated import.',
    ],
    'templates_intro': (
        'Twelve template / page families were identified across the 439 pages. The site uses '
        'four different page "shells" (standard, corporate, minimal microsite, legacy /en/), '
        'which is the main driver of complexity.'
    ),
    'templates': [
        ['Template', 'Complexity', 'Reasoning', 'Reference URLs'],
        ['Homepage', 'High',
         'Hero/banner, segment links, notice ribbon, news feed, promo cards. One-off layout.',
         'https://www.sumitclub.jp/ja/index.html'],
        ['Category / Service / How-to', 'Medium',
         'Standard /ja/ shell content: usage, travel, point, insurance, entertainment, '
         'campaign. Page heading, sub-nav, image+text, step blocks, tables, CTA (175 pages).',
         'https://www.sumitclub.jp/ja/usage/clubonline.html\n'
         'https://www.sumitclub.jp/ja/travel/airport.html'],
        ['Article / News detail', 'Low',
         'Dated notice articles incl. /en/ locale: date, title, rich text (109 pages).',
         'https://www.sumitclub.jp/en/notice/20190301.html'],
        ['Corporate info', 'Medium',
         'Corporate-site shell (separate header/nav/footer): company overview, philosophy, '
         'message, locations, service/SDGs pages (60 pages).',
         'https://www.sumitclub.jp/ja/corporate_site/company/overview.html'],
        ['Login / Member system page', 'High',
         'Club Online member-system screens (WA######## ids). System-generated, '
         'authenticated, dynamic — not editorial content (40 pages).',
         'https://www.sumitclub.jp/ja/loginPage/WA20501040002A.html'],
        ['Product / Card detail', 'Medium',
         'Card-spec layout: hero, feature/fee tables, related-cards carousel, apply CTA to '
         'external entry form (20 pages).',
         'https://www.sumitclub.jp/ja/cardlineup/deltagold.html'],
        ['Application / Verification microsite', 'High',
         'Minimal-chrome shells for identity verification / application help (honninkakunin, '
         'ccol bumper pages, id-doc). Often launch external apps via JS (15 pages).',
         'https://www.sumitclub.jp/ja/honninkakunin.html'],
        ['Legal / Policy text', 'Low',
         'Long-form text + tables: privacy, small print, terms (7 pages).',
         'https://www.sumitclub.jp/ja/privacy.html'],
        ['Error / System page', 'Low',
         'HTTP error pages (400/401/403/404/500/501/503). Platform-level, likely replaced by '
         'EDS error handling (7 pages).',
         'https://www.sumitclub.jp/ja/404.html'],
        ['Contact / Form', 'Medium',
         'Contact directory + form-gateway page that launches a JS inquiry widget (2 pages).',
         'https://www.sumitclub.jp/ja/contact.html\n'
         'https://www.sumitclub.jp/ja/contact_form.html'],
        ['Listing / Index', 'Medium',
         'Notice index and card-services comparison list with grouping/load-more (2 pages).',
         'https://www.sumitclub.jp/ja/notice.html\n'
         'https://www.sumitclub.jp/en/cardlineup/card_services_list.html'],
        ['Corporate news list', 'Low',
         'Dated corporate news listing (1 page).',
         'https://www.sumitclub.jp/ja/corporate_site/news.html'],
        ['Sitemap', 'Low', 'Single hierarchical link-list page.',
         'https://www.sumitclub.jp/ja/sitemap.html'],
    ],
    'templates_note': (
        'Note: four distinct page shells exist (standard /ja/, corporate-site, minimal '
        'microsite, legacy /en/). Consolidating these into one EDS header/footer + section '
        'styling is recommended and is the single biggest design decision for this site. '
        'A migration-time hazard was observed: Adobe Target/Analytics occasionally redirects '
        'sumitclub /ja/ pages to diners.co.jp during navigation — this must be disabled for '
        'the migrated estate.'
    ),
    'blocks_intro': (
        'The following 25 reusable blocks were identified. Per the brief, visually different '
        'layouts sharing one content model are captured as variants rather than new blocks. '
        'Header/footer have four shell variants today that should consolidate to one.'
    ),
    'blocks': [
        COMMON_BLOCKS_HEADER,
        ['Global Header / Nav', '4', 'High',
         'Four variants today: standard /ja/, corporate (dropdown), minimal (logo-only), '
         'legacy /en/ (hamburger). Recommend consolidating to one EDS nav.'],
        ['Global Footer', '4', 'Medium',
         'Matching four variants (standard, corporate mega-footer, slim microsite, legacy '
         '/en/ mega-links). Consolidate to one.'],
        ['Breadcrumb', '2', 'Low',
         'Path nav. Variant: corporate "現在位置" breadcrumb.'],
        ['Section Sub-Nav', '1', 'Low', 'In-section secondary navigation.'],
        ['Hero / Page Heading', '2', 'Low',
         'Title + optional image. Variant: card-image hero on product pages.'],
        ['Card Grid / Listing', '2', 'Medium',
         'Linked card grid. Variants: service cards, "FAQ + apply" promo cards.'],
        ['Related-Cards Carousel', '1', 'Medium',
         'Prev/next + tablist carousel of related cards on product pages.'],
        ['News List', '2', 'Medium',
         'Dated list with category badge + load-more. Variants: year-grouped accordion list, '
         'flat dl list.'],
        ['Rich Text', '1', 'Low', 'Body copy with headings, lists, links.'],
        ['Data Table', '2', 'Low',
         'Spec/fee tables. Variant: service-comparison table.'],
        ['Image + Caption', '1', 'Low', 'Media with caption.'],
        ['Image + Text', '2', 'Low', 'Media-and-copy. Variants: image-left, image-right.'],
        ['Feature / Icon Grid', '2', 'Low',
         'Icon+label grid. Variants: feature icons, payment-method icons.'],
        ['Step / Process Block', '1', 'Low', 'Numbered STEP sequence.'],
        ['Accordion / Toggle', '1', 'Low', 'Expand/collapse sections (incl. news grouping).'],
        ['Tabs', '1', 'Medium', 'Category tabs / tablist.'],
        ['CTA Button', '2', 'Low', 'Primary CTA. Variants: internal, external/apply.'],
        ['Link List', '2', 'Low', 'Grouped link lists. Variant: sitemap tree.'],
        ['Phone / Contact Directory', '1', 'Low', 'tel: number directory block.'],
        ['Notice / Ribbon', '1', 'Low', 'Highlighted notice strip.'],
        ['Category Tag / Badge', '1', 'Low', 'Notice category labels.'],
        ['Search Box', '1', 'Medium',
         'Header search (legacy /en/ posts to search.diners.co.jp). External engine.'],
        ['Corporate Footer Block', '1', 'Medium',
         'Corporate mega-footer with company link groups (corporate shell only).'],
        ['Form Gateway / Widget Launcher', '1', 'Medium',
         'Button that injects/launches external inquiry or insurer application widget.'],
        ['Error Page Block', '1', 'Low', 'System error message layout.'],
    ],
    'reconcile_intro': (
        'The team\'s Excel block/integration sheet (Sheet1 = blocks, Sheet2 = integrations) '
        'was not available to this analysis at the time of writing. The table below lists the '
        'live-site findings for diffing against the team list. Please confirm the column; '
        'items marked "Confirm" are additions or consolidations this analysis recommends.'
    ),
    'reconcile': [
        ['Recommendation', 'Block(s)', 'Rationale'],
        ['Consolidate as variants',
         'Header (4 shells → 1); Footer (4 → 1); News List (accordion / flat); '
         'Card Grid (service / promo); Data Table (spec / comparison); Image+Text (L/R); '
         'Feature Grid (feature / payment icons)',
         'Same content model, different visual layout — one block + variants. The 4-shell '
         'header/footer consolidation is the biggest single saving.'],
        ['Confirm / likely additions',
         'Section Sub-Nav, Related-Cards Carousel, Step/Process, Phone Directory, '
         'Form Gateway launcher, Error Page block, Corporate Footer',
         'Present on the live site across the different shells; verify against the team list.'],
        ['Treat as integration / out of scope',
         'Login / Member system pages (40), Verification microsites, Search Box, '
         'Insurer application widget',
         'Functional/system pages and external gateways — scope as integrations or exclude, '
         'not as editorial content blocks.'],
    ],
    'counts_intro': (
        'Page counts below are derived by classifying all 439 sitemap URLs (across /ja/, '
        '/en/ and corporate areas) into the twelve template families from Section 2.'
    ),
    'counts': [
        ['Template', 'Pages', '% of site', 'Migration approach', 'Rationale'],
        ['Category / Service / How-to', '175', '39.9%', 'Assisted',
         'Standard shell, per-section layout variation; bulk import + cleanup.'],
        ['Article / News detail', '109', '24.8%', 'Automated',
         'Standardized dated notices (incl. /en/).'],
        ['Corporate info', '60', '13.7%', 'Assisted',
         'Separate corporate shell; bulk import after shell mapping.'],
        ['Login / Member system page', '40', '9.1%', 'Manual / Exclude',
         'System-generated authenticated pages — re-point to Club Online, likely out of scope.'],
        ['Product / Card detail', '20', '4.6%', 'Manual',
         'Marketing-critical card pages; rebuild for fidelity + external apply CTA.'],
        ['Application / Verification microsite', '15', '3.4%', 'Manual',
         'Minimal-chrome shells launching external apps; case-by-case.'],
        ['Legal / Policy text', '7', '1.6%', 'Automated', 'Plain long-form text + tables.'],
        ['Error / System page', '7', '1.6%', 'Manual / Exclude',
         'Platform error pages — replace with EDS error handling.'],
        ['Listing / Index', '2', '0.5%', 'Manual', 'Grouping / load-more needs EDS index + query.'],
        ['Contact / Form', '2', '0.5%', 'Manual', 'External form widget wiring.'],
        ['Corporate news list', '1', '0.2%', 'Automated', 'Dated list page.'],
        ['Sitemap', '1', '0.2%', 'Automated', 'Single link-list page.'],
        ['TOTAL', '439', '100%', '—', '—'],
    ],
    'integrations_intro': (
        'The site embeds several third-party / external systems, several of which are shared '
        'with diners.co.jp. Most should be re-pointed or re-embedded rather than rebuilt.'
    ),
    'integrations': [
        ['Integration', 'Purpose', 'Migration handling'],
        ['Club Online (sumitclub.jp/JPCRD/col)', 'Member login / self-service', 'External member system — link out; the 40 loginPage URLs likely stay on it.'],
        ['Helpfeel', 'FAQ / help', 'External — outbound link / embed.'],
        ['Card application LPs (entry_form/lp)', 'New-card applications', 'External landing pages — link out.'],
        ['Inquiry form widget', 'Contact form', 'JS-injected widget — re-embed or replace with EDS form.'],
        ['Third-party insurer app (winop)', 'Golf/insurance application', 'External — link out.'],
        ['Remote Operator (sumitclub.rmop.jp)', 'Co-browse support', 'External — link out.'],
        ['Adobe Analytics / Target + adobe_mc', 'Analytics & personalization, cross-domain ID', 'Re-tag; DISABLE the Target rule that redirects /ja/ pages to diners.co.jp.'],
        ['Meta Pixel / im-apps / logly / doga.cm', 'Advertising & remarketing', 'Re-add required tags only.'],
        ['Online Mall (sumitclubonlinemall.jp)', 'Points shopping mall', 'External — link out.'],
        ['Site Search (search.diners.co.jp)', 'On-site search (legacy /en/)', 'Re-point or replace with EDS/Edge search.'],
    ],
    'estimates_intro': (
        'Estimates assume reuse of the EDS import accelerator and, ideally, shared block '
        'development with diners.co.jp (same platform, shared header/footer). The page volume '
        'is modest (439) but shell fragmentation and the high share of system/functional '
        'pages push up the manual portion.'
    ),
    'effort': [
        ['Work-stream', 'Effort (person-days)', 'Notes'],
        ['Discovery & template design', '5–6', 'Confirm 12 families + 4-shell consolidation.'],
        ['Block development (25 blocks, with variants)', '16–20',
         'Lower if shared with diners.co.jp; 4-shell header/footer consolidation is key.'],
        ['Import pipeline (parsers + transformers)', '8–10',
         'Parsers for usage/how-to, notice, corporate, product.'],
        ['Automated content migration (≈300 pages)', '5–7', 'Bulk runs + reruns.'],
        ['Manual / assisted pages (≈139 pages)', '12–16',
         'Card pages, microsites, corporate cleanup, listings, login decisions.'],
        ['Integrations re-wiring', '7–9', 'Search, forms, video, analytics, login links, redirect fix.'],
        ['QA & Testing', '9–11', 'Visual QA, link checks, responsive, bilingual (/ja/ + /en/).'],
        ['PM / coordination / UAT support', '4–6', 'Across the engagement.'],
        ['TOTAL', '60–74', '≈ 9–11 weeks elapsed with a 2–3 person team.'],
    ],
    'content_effort': [
        ['Template', 'Pages', 'Approach', 'Effort'],
        ['Category / Service / How-to', '175', 'Assisted', '4–5 days (import + cleanup)'],
        ['Article / News detail', '109', 'Automated', '1–1.5 days'],
        ['Corporate info', '60', 'Assisted', '2.5–3 days'],
        ['Login / Member system page', '40', 'Manual / Exclude', '1–2 days (decision + re-point)'],
        ['Product / Card detail', '20', 'Manual', '2.5–3 days'],
        ['Application / Verification microsite', '15', 'Manual', '2–3 days'],
        ['Legal / Policy text', '7', 'Automated', '0.5 day'],
        ['Error / System page', '7', 'Manual / Exclude', '0.5 day'],
        ['Listing / Index', '2', 'Manual', '1.5–2 days'],
        ['Contact / Form', '2', 'Manual', '1–2 days'],
        ['Corporate news list', '1', 'Automated', '0.25 day'],
        ['Sitemap', '1', 'Automated', '0.25 day'],
    ],
    'timeline': [
        ['Phase', 'Duration', 'Key outputs'],
        ['Phase 1 — Discovery & design', 'Weeks 1–2', 'Template catalog, shell consolidation plan, block models.'],
        ['Phase 2 — Block & pipeline build', 'Weeks 2–5', 'Blocks + import parsers/transformers.'],
        ['Phase 3 — Bulk migration', 'Weeks 5–7', 'Automated import of ~300 standardized pages.'],
        ['Phase 4 — Manual pages & integrations', 'Weeks 7–10', 'Card pages, microsites, corporate, integrations, redirect fix.'],
        ['Phase 5 — QA, UAT & launch', 'Weeks 9–11', 'Bilingual QA, fixes, go-live.'],
    ],
    'cost': [
        ['Scenario', 'Effort (person-days)', 'Indicative range*'],
        ['Low (optimistic)', '60', '60 × blended day-rate'],
        ['High (conservative)', '74', '74 × blended day-rate'],
    ],
    'cost_note': (
        '*Cost is expressed in person-days; apply your blended daily rate to convert to '
        'currency. As an illustration, at a representative offshore/near-shore blended rate '
        'of USD 500–700/day, the engagement falls roughly in the USD 30k–52k range. If '
        'migrated together with diners.co.jp, shared block/header/footer development could '
        'reduce combined effort by an estimated 15–20%. Replace with your contracted rate.'
    ),
    'assumptions': [
        'Page counts are taken from the live sitemaps (/sitemap.xml and /ja/Sitemap.xml) on '
        'the analysis date; 439 published pages across /ja/, /en/ and corporate areas.',
        'Header/footer and the Club Online member system are shared with www.diners.co.jp; '
        'migrating both sites together amortizes shared development.',
        'The 40 login/member-system pages and 7 system error pages are treated as functional/'
        'out-of-scope and re-pointed or replaced rather than migrated as content.',
        'The Adobe Target rule that redirects sumitclub /ja/ pages to diners.co.jp must be '
        'disabled for the migrated estate.',
        'External systems (Club Online, Helpfeel, application LPs, insurer app, search) remain '
        'on their current platforms and are linked/embedded, not rebuilt.',
        'The team\'s Excel block/integration list was not available at writing time; Section '
        '3.1 should be reconciled against it before finalizing scope.',
        'Estimates assume a 2–3 engineer + 1 QA team and reuse of the EDS import accelerator; '
        'they exclude content rewriting, new design, translation, and post-launch support.',
    ],
}


if __name__ == '__main__':
    build_report(diners, '/workspace/current/Diners_Club_Japan_EDS_Migration_Analysis.docx')
    build_report(sumitclub, '/workspace/current/TrustClub_Sumitclub_EDS_Migration_Analysis.docx')
